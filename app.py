import os, io, sqlite3, hashlib, hmac, secrets, uuid, base64, json
from functools import wraps
from datetime import datetime

import qrcode, requests as http_requests
from flask import Flask, request, jsonify, session, send_from_directory, redirect, after_this_request
from urllib.parse import urlencode
from werkzeug.middleware.proxy_fix import ProxyFix

# Load .env file if present
_env_path = os.path.join(os.path.dirname(__file__), '.env')
if os.path.exists(_env_path):
    with open(_env_path) as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith('#') and '=' in _line:
                _k, _v = _line.split('=', 1)
                _key, _val = _k.strip(), _v.strip()
                if _val or _key not in os.environ:
                    os.environ[_key] = _val

app = Flask(__name__, static_folder=None)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)
app.secret_key = os.environ.get('SECRET_KEY') or secrets.token_hex(32)
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Railway: смонтируй Volume на /data, задай DATABASE_PATH=/data/database.db
_vol = os.environ.get('RAILWAY_VOLUME_MOUNT_PATH')
DB_PATH = os.environ.get('DATABASE_PATH') or (
    os.path.join(_vol, 'database.db') if _vol else
    os.path.join(os.path.dirname(__file__), 'database.db')
)
OPENAI_KEY = os.environ.get('OPENAI_API_KEY', '')
GEMINI_KEY = os.environ.get('GEMINI_API_KEY', '')
GROQ_KEY = os.environ.get('GROQ_API_KEY', '')
OPENROUTER_KEY = os.environ.get('OPENROUTER_API_KEY', '')
DEEPSEEK_KEY = os.environ.get('DEEPSEEK_API_KEY', '')
CEREBRAS_KEY = os.environ.get('CEREBRAS_API_KEY', '')
GIGACHAT_AUTH_KEY = os.environ.get('GIGACHAT_AUTH_KEY', '')
SMTP_EMAIL = os.environ.get('SMTP_EMAIL', '')
SMTP_PASSWORD = os.environ.get('SMTP_PASSWORD', '')
RESEND_API_KEY = (os.environ.get('RESEND_API_KEY') or '').strip()
MAILGUN_API_KEY = (os.environ.get('MAILGUN_API_KEY') or '').strip()
MAILGUN_DOMAIN = (os.environ.get('MAILGUN_DOMAIN') or '').strip()
MAILGUN_FROM = (os.environ.get('MAILGUN_FROM') or '').strip() or (('noreply@' + MAILGUN_DOMAIN) if MAILGUN_DOMAIN else '')

YANDEX_CLIENT_ID = os.environ.get('YANDEX_CLIENT_ID', '')
YANDEX_CLIENT_SECRET = os.environ.get('YANDEX_CLIENT_SECRET', '')
VK_CLIENT_ID = os.environ.get('VK_CLIENT_ID', '')
VK_CLIENT_SECRET = os.environ.get('VK_CLIENT_SECRET', '')
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID', '')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET', '')
TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', '')
TELEGRAM_BOT_USERNAME = os.environ.get('TELEGRAM_BOT_USERNAME', '')
ADMIN_EMAILS = [e.strip().lower() for e in os.environ.get('ADMIN_EMAILS', '').split(',') if e.strip()]
OAUTH_BASE_URL = (os.environ.get('OAUTH_BASE_URL') or '').strip().rstrip('/')


def _oauth_origin():
    """Base URL for OAuth redirect_uri (tunnel or current request)."""
    return OAUTH_BASE_URL if OAUTH_BASE_URL else request.host_url.rstrip('/')


# ─── БД ──────────────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn

def init_db():
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            first_name TEXT NOT NULL DEFAULT '',
            last_name TEXT NOT NULL DEFAULT '',
            birthdate TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS polls (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            slug TEXT UNIQUE NOT NULL,
            multiple_choice INTEGER DEFAULT 0,
            is_active INTEGER DEFAULT 1,
            auth_only INTEGER DEFAULT 0,
            show_results TEXT DEFAULT 'always',
            deadline TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS poll_options (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            poll_id INTEGER NOT NULL,
            text TEXT NOT NULL,
            sort_order INTEGER DEFAULT 0,
            FOREIGN KEY (poll_id) REFERENCES polls(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS poll_votes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            poll_id INTEGER NOT NULL,
            option_id INTEGER NOT NULL,
            user_id INTEGER,
            fingerprint TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (poll_id) REFERENCES polls(id) ON DELETE CASCADE,
            FOREIGN KEY (option_id) REFERENCES poll_options(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS decisions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            slug TEXT UNIQUE NOT NULL,
            auth_only INTEGER DEFAULT 0,
            deadline TEXT,
            is_active INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS decision_alternatives (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            decision_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            sort_order INTEGER DEFAULT 0,
            FOREIGN KEY (decision_id) REFERENCES decisions(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS decision_criteria (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            decision_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            sort_order INTEGER DEFAULT 0,
            FOREIGN KEY (decision_id) REFERENCES decisions(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS decision_responses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            decision_id INTEGER NOT NULL,
            user_id INTEGER,
            fingerprint TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (decision_id) REFERENCES decisions(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS decision_scores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            response_id INTEGER NOT NULL,
            alternative_id INTEGER NOT NULL,
            criterion_id INTEGER NOT NULL,
            score INTEGER NOT NULL DEFAULT 0,
            FOREIGN KEY (response_id) REFERENCES decision_responses(id) ON DELETE CASCADE,
            FOREIGN KEY (alternative_id) REFERENCES decision_alternatives(id) ON DELETE CASCADE,
            FOREIGN KEY (criterion_id) REFERENCES decision_criteria(id) ON DELETE CASCADE
        );
    """)
    conn.commit()
    conn.close()

init_db()

def _migrate():
    conn = get_db()
    cols_polls = [r['name'] for r in conn.execute("PRAGMA table_info(polls)").fetchall()]
    if 'show_voters' not in cols_polls:
        conn.execute("ALTER TABLE polls ADD COLUMN show_voters INTEGER DEFAULT 0")
    if 'anonymous' not in cols_polls:
        conn.execute("ALTER TABLE polls ADD COLUMN anonymous INTEGER DEFAULT 0")
    if 'max_votes' not in cols_polls:
        conn.execute("ALTER TABLE polls ADD COLUMN max_votes INTEGER DEFAULT 0")
    cols_decs = [r['name'] for r in conn.execute("PRAGMA table_info(decisions)").fetchall()]
    if 'show_respondents' not in cols_decs:
        conn.execute("ALTER TABLE decisions ADD COLUMN show_respondents INTEGER DEFAULT 0")
    if 'anonymous' not in cols_decs:
        conn.execute("ALTER TABLE decisions ADD COLUMN anonymous INTEGER DEFAULT 0")
    if 'show_results' not in cols_decs:
        conn.execute("ALTER TABLE decisions ADD COLUMN show_results TEXT DEFAULT 'always'")
    if 'is_public' not in cols_decs:
        conn.execute("ALTER TABLE decisions ADD COLUMN is_public INTEGER DEFAULT 0")
    if 'scale_max' not in cols_decs:
        conn.execute("ALTER TABLE decisions ADD COLUMN scale_max INTEGER DEFAULT 5")
    if 'is_public' not in cols_polls:
        conn.execute("ALTER TABLE polls ADD COLUMN is_public INTEGER DEFAULT 0")
    cols_users = [r['name'] for r in conn.execute("PRAGMA table_info(users)").fetchall()]
    if 'oauth_provider' not in cols_users:
        conn.execute("ALTER TABLE users ADD COLUMN oauth_provider TEXT DEFAULT ''")
    if 'oauth_id' not in cols_users:
        conn.execute("ALTER TABLE users ADD COLUMN oauth_id TEXT DEFAULT ''")
    if 'warnings' not in cols_users:
        conn.execute("ALTER TABLE users ADD COLUMN warnings INTEGER DEFAULT 0")
    if 'is_blocked' not in cols_users:
        conn.execute("ALTER TABLE users ADD COLUMN is_blocked INTEGER DEFAULT 0")
    conn.execute("""CREATE TABLE IF NOT EXISTS moderation_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        content_type TEXT NOT NULL DEFAULT 'poll',
        poll_title TEXT DEFAULT '',
        poll_description TEXT DEFAULT '',
        reason TEXT DEFAULT '',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS poll_reactions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        poll_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        type TEXT NOT NULL CHECK(type IN ('like','dislike')),
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (poll_id) REFERENCES polls(id) ON DELETE CASCADE,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        UNIQUE(poll_id, user_id)
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS poll_comments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        poll_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        text TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (poll_id) REFERENCES polls(id) ON DELETE CASCADE,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS decision_comments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        decision_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        text TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (decision_id) REFERENCES decisions(id) ON DELETE CASCADE,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS decision_reactions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        decision_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        type TEXT NOT NULL CHECK(type IN ('like','dislike')),
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (decision_id) REFERENCES decisions(id) ON DELETE CASCADE,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        UNIQUE(decision_id, user_id)
    )""")
    conn.commit()
    conn.close()

_migrate()


# ─── Утилиты ─────────────────────────────────────────────────
def hash_password(pw):
    salt = secrets.token_hex(16)
    return salt + ":" + hashlib.sha256((salt + pw).encode()).hexdigest()

def check_password(pw, stored):
    salt, h = stored.split(":")
    return hashlib.sha256((salt + pw).encode()).hexdigest() == h

def user_dict(row):
    d = {"id": row["id"], "email": row["email"], "firstName": row["first_name"],
         "lastName": row["last_name"], "birthdate": row["birthdate"], "createdAt": row["created_at"]}
    try:
        d["isBlocked"] = bool(row["is_blocked"])
        d["warnings"] = row["warnings"]
    except (IndexError, KeyError):
        pass
    try:
        d["isAdmin"] = ((row["email"] or "").strip().lower() in ADMIN_EMAILS)
    except (IndexError, KeyError):
        d["isAdmin"] = False
    return d

def get_fingerprint():
    return hashlib.md5((request.remote_addr or '' + request.headers.get('User-Agent', '')).encode()).hexdigest()

def check_deadline(deadline_str):
    if not deadline_str:
        return True
    try:
        dl = datetime.fromisoformat(deadline_str)
        return datetime.now() < dl
    except:
        return True


def _call_llm(base_url, api_key, model, system_prompt, user_prompt, temperature, max_tokens, extra_headers=None):
    """Generic OpenAI-compatible call with 1 retry. Returns text or None."""
    import time as _time
    from openai import OpenAI
    kwargs = {"base_url": base_url, "api_key": api_key, "timeout": 60.0}
    if extra_headers:
        kwargs["default_headers"] = extra_headers
    client = OpenAI(**kwargs)
    for attempt in range(2):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=temperature, max_tokens=max_tokens
            )
            txt = resp.choices[0].message.content
            if txt:
                return txt.strip()
        except Exception as e:
            print(f"[LLM {model} attempt {attempt+1}] {e}")
            if attempt == 0:
                _time.sleep(2)
    return None


_OPENROUTER_MODELS = [
    "meta-llama/llama-3.3-70b-instruct:free",
    "qwen/qwen3-32b:free",
    "google/gemma-3-27b-it:free",
    "mistralai/mistral-small-3.1-24b-instruct:free",
    "microsoft/phi-4:free",
]


def _call_gigachat(system_prompt, user_prompt, temperature=0.7, max_tokens=700):
    """Call Sber GigaChat API. Returns (text, error_message)."""
    if not GIGACHAT_AUTH_KEY:
        return None, "GIGACHAT_AUTH_KEY не задан в .env"
    try:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        tok_resp = http_requests.post(
            'https://ngw.devices.sberbank.ru:9443/api/v2/oauth',
            headers={'Content-Type': 'application/x-www-form-urlencoded',
                     'Accept': 'application/json',
                     'RqUID': str(uuid.uuid4()),
                     'Authorization': 'Basic ' + GIGACHAT_AUTH_KEY},
            data={'scope': 'GIGACHAT_API_PERS'},
            verify=False, timeout=20
        )
        try:
            tok_data = tok_resp.json()
        except Exception:
            return None, "GigaChat: неверный ответ при получении токена (%s)" % (tok_resp.text[:200] if tok_resp.text else tok_resp.status_code)
        access_token = tok_data.get('access_token', '')
        if not access_token:
            err = tok_data.get('error_description') or tok_data.get('error') or tok_resp.text or str(tok_resp.status_code)
            return None, "GigaChat токен: %s" % (err[:150] if isinstance(err, str) else err)
        resp = http_requests.post(
            'https://gigachat.devices.sberbank.ru/api/v1/chat/completions',
            headers={'Authorization': 'Bearer ' + access_token,
                     'Content-Type': 'application/json',
                     'Accept': 'application/json'},
            json={'model': 'GigaChat',
                  'messages': [
                      {'role': 'system', 'content': system_prompt},
                      {'role': 'user', 'content': user_prompt}
                  ],
                  'temperature': temperature,
                  'max_tokens': max_tokens},
            verify=False, timeout=90
        )
        try:
            data = resp.json()
        except Exception:
            return None, "GigaChat: неверный ответ (%s)" % (resp.text[:200] if resp.text else resp.status_code)
        choices = data.get('choices') or []
        if choices:
            msg = choices[0].get('message') or choices[0]
            txt = (msg.get('content') or msg.get('text') or '').strip()
            if txt:
                return txt, None
        err = data.get('error', {}) or {}
        err_msg = err.get('message') or err.get('error') or resp.text or str(resp.status_code)
        return None, "GigaChat: %s" % (err_msg[:150] if isinstance(err_msg, str) else err_msg)
    except Exception as e:
        err_str = str(e)
        print(f'[GigaChat error] {e}')
        return None, "GigaChat: %s" % (err_str[:150] if len(err_str) > 150 else err_str)


def _call_ai(system_prompt, user_prompt, temperature=0.7, max_tokens=700):
    """Try GigaChat -> OpenRouter -> Cerebras -> DeepSeek -> Groq -> Gemini -> OpenAI. Returns (text, source, last_error)."""
    last_error = None
    if GIGACHAT_AUTH_KEY:
        text, err = _call_gigachat(system_prompt, user_prompt, temperature, max_tokens)
        if text:
            return text, "gigachat", None
        last_error = err or last_error
    if OPENROUTER_KEY:
        or_headers = {"HTTP-Referer": "http://localhost:5000", "X-Title": "Decision Service"}
        for model in _OPENROUTER_MODELS:
            text = _call_llm("https://openrouter.ai/api/v1", OPENROUTER_KEY,
                              model, system_prompt, user_prompt, temperature, max_tokens, or_headers)
            if text:
                return text, "llama", None
    if CEREBRAS_KEY:
        text = _call_llm("https://api.cerebras.ai/v1", CEREBRAS_KEY,
                          "llama3.1-8b", system_prompt, user_prompt, temperature, max_tokens)
        if text:
            return text, "cerebras", None
    if DEEPSEEK_KEY:
        text = _call_llm("https://api.deepseek.com/v1", DEEPSEEK_KEY,
                          "deepseek-chat", system_prompt, user_prompt, temperature, max_tokens)
        if text:
            return text, "deepseek", None
    if GROQ_KEY:
        text = _call_llm("https://api.groq.com/openai/v1", GROQ_KEY,
                          "llama-3.3-70b-versatile", system_prompt, user_prompt, temperature, max_tokens)
        if text:
            return text, "groq", None
    if GEMINI_KEY:
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=GEMINI_KEY)
            resp = client.models.generate_content(
                model="gemini-2.5-flash", contents=user_prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt, temperature=temperature, max_output_tokens=max_tokens))
            if resp.text:
                return resp.text.strip(), "gemini", None
        except Exception as e:
            last_error = "Gemini: %s" % str(e)[:120]
    if OPENAI_KEY:
        text = _call_llm("https://api.openai.com/v1", OPENAI_KEY,
                          "gpt-3.5-turbo", system_prompt, user_prompt, temperature, max_tokens)
        if text:
            return text, "openai", None
    return None, None, last_error or "Нет ключей ИИ в .env (GIGACHAT_AUTH_KEY, OPENROUTER_KEY и др.)"


_MODERATION_PROMPT = (
    "Ты автоматический фильтр контента. Твоя задача — классифицировать текст. "
    "Правила: запрещены мат, оскорбления, угрозы, призывы к насилию, разжигание ненависти, порнография, спам. "
    "ВАЖНО: Ты ОБЯЗАН ответить РОВНО одним словом. Никаких пояснений. "
    "Если текст допустим — ответь: OK\n"
    "Если текст нарушает правила — ответь: REJECT"
)

import re as _re
_BAD_WORDS = _re.compile(
    r'(бля[дт]ь?|бля|[хx][уy][йеёия]|пизд|ебан|ёбан|еба[тл]|сук[аи]|мудак|мудил'
    r'|дерьм|гандон|пидор|пидар|шлюх|хуес|залуп|манд[аоуе]|долбоёб|долбоеб'
    r'|уёбок|уебок|уёбищ|наху[йея]|нахер|похуй|отпизд|выеб|трахн'
    r'|убить|убей|сдохн|зарежу|застрел|взорв[аеу]|повес)',
    _re.IGNORECASE | _re.UNICODE
)

def _moderate_content(title, description='', options=None):
    """AI content moderation. Returns (ok: bool, reason: str|None)."""
    parts = []
    if title:
        parts.append(f"Заголовок: {title}")
    if description:
        parts.append(f"Описание: {description}")
    if options:
        opts_text = ', '.join(o if isinstance(o, str) else o.get('text', '') for o in options)
        parts.append(f"Варианты: {opts_text}")
    full_text = '\n'.join(parts)

    bad = _BAD_WORDS.search(full_text)
    if bad:
        reason = 'Нецензурная лексика или призыв к насилию'
        print(f'[Moderation] LOCAL REJECT: matched "{bad.group()}" in text')
        return False, reason

    try:
        print(f'[Moderation] AI checking: {full_text[:100]}')
        result, source, _ = _call_ai(_MODERATION_PROMPT, full_text, temperature=0.1, max_tokens=30)
        print(f'[Moderation] AI response: {result!r} (source={source})')
        if result is None:
            print('[Moderation] AI недоступен — пропуск проверки (действует только локальный фильтр)')
            return True, None
        r = result.strip().upper().replace('\u041e', 'O').replace('\u041a', 'K')  # кириллица О,К → латиница
        if r.startswith('OK'):
            return True, None
        if 'REJECT' in r:
            return False, 'Контент нарушает правила сообщества'
        # Убрано «не могу» — ИИ часто пишет это при неуверенности, что давало ложные срабатывания
        if any(w in result.lower() for w in ['недопустим', 'нарушает', 'запрещён', 'запрещен', 'отклон',
                                              'нецензурн', 'насили', 'не поддержив']):
            return False, 'Контент отклонён модерацией'
        return True, None
    except Exception as e:
        print(f'[Moderation error] {e}')
        return True, None

def _apply_warning(user_id, content_type, title, description, reason):
    """Increment warning, block if >= 3, log to moderation_log. Returns (warnings, is_blocked)."""
    conn = get_db()
    conn.execute("UPDATE users SET warnings = warnings + 1 WHERE id=?", (user_id,))
    user = conn.execute("SELECT warnings FROM users WHERE id=?", (user_id,)).fetchone()
    w = user['warnings']
    blocked = w >= 3
    if blocked:
        conn.execute("UPDATE users SET is_blocked=1 WHERE id=?", (user_id,))
    conn.execute(
        "INSERT INTO moderation_log (user_id, content_type, poll_title, poll_description, reason) VALUES (?,?,?,?,?)",
        (user_id, content_type, title or '', description or '', reason or ''))
    conn.commit()
    conn.close()
    return w, blocked


def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({"error": "Необходима авторизация"}), 401
        return f(*args, **kwargs)
    return wrapper

def block_if_banned(f):
    """Like login_required but also rejects blocked users for write endpoints."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({"error": "Необходима авторизация"}), 401
        conn = get_db()
        user = conn.execute("SELECT is_blocked FROM users WHERE id=?", (session['user_id'],)).fetchone()
        conn.close()
        if user and user['is_blocked']:
            return jsonify({"error": "Аккаунт заблокирован", "blocked": True}), 403
        return f(*args, **kwargs)
    return wrapper

def admin_required(f):
    """Require logged-in user whose email is in ADMIN_EMAILS."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({"error": "Необходима авторизация"}), 401
        if not ADMIN_EMAILS:
            return jsonify({"error": "Админ-панель не настроена"}), 403
        conn = get_db()
        user = conn.execute("SELECT email FROM users WHERE id=?", (session['user_id'],)).fetchone()
        conn.close()
        if not user or (user['email'] or '').strip().lower() not in ADMIN_EMAILS:
            return jsonify({"error": "Доступ запрещён"}), 403
        return f(*args, **kwargs)
    return wrapper


@app.route('/', methods=['GET', 'POST'])
def serve_index():
    return send_from_directory('.', 'index.html')


# ═══════════════════════════════════════════════════════════════
#  AUTH
# ═══════════════════════════════════════════════════════════════
@app.route('/api/register', methods=['GET', 'POST'])
def register():
    if request.method == 'GET':
        return redirect('/register.html')
    d = request.get_json()
    email = (d.get('email') or '').strip().lower()
    pw = d.get('password') or ''
    fn = (d.get('firstName') or '').strip()
    ln = (d.get('lastName') or '').strip()
    bd = (d.get('birthdate') or '').strip()
    if not email or not pw: return jsonify({"error": "Email и пароль обязательны"}), 400
    if len(pw) < 6: return jsonify({"error": "Пароль — не менее 6 символов"}), 400
    if not fn or not ln: return jsonify({"error": "Имя и фамилия обязательны"}), 400
    if not d.get('agreePd') or not d.get('agreeTerms'): return jsonify({"error": "Необходимо принять все соглашения"}), 400
    oauth = session.pop('oauth_pending', None)
    oauth_provider = oauth['provider'] if oauth else ''
    oauth_id = oauth['oauth_id'] if oauth else ''
    conn = get_db()
    try:
        conn.execute("INSERT INTO users (email,password_hash,first_name,last_name,birthdate,oauth_provider,oauth_id) VALUES (?,?,?,?,?,?,?)",
                     (email, hash_password(pw), fn, ln, bd or None, oauth_provider, oauth_id))
        conn.commit()
        user = conn.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
        session['user_id'] = user['id']
        return jsonify({"ok": True, "user": user_dict(user)}), 201
    except sqlite3.IntegrityError:
        return jsonify({"error": "Пользователь с таким email уже существует"}), 409
    finally:
        conn.close()

@app.route('/api/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return redirect('/login.html')
    d = request.get_json()
    email = (d.get('email') or '').strip().lower()
    pw = d.get('password') or ''
    if not email or not pw: return jsonify({"error": "Email и пароль обязательны"}), 400
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
    conn.close()
    if not user or not check_password(pw, user['password_hash']): return jsonify({"error": "Неверный email или пароль"}), 401
    session['user_id'] = user['id']
    return jsonify({"ok": True, "user": user_dict(user)})

@app.route('/api/logout', methods=['GET', 'POST'])
def logout():
    session.clear()
    if request.method == 'GET':
        return redirect('/')
    return jsonify({"ok": True})

_reset_codes = {}

def _send_email(to, subject, body):
    """Send via Resend (без домена!) -> Mailgun -> Gmail SMTP. Returns (success: bool, error_message: str|None)."""
    from email.utils import formataddr
    from email.mime.text import MIMEText
    from email.header import Header

    # Resend — без домена, на любой email. Бесплатно 100 писем/день.
    if RESEND_API_KEY:
        try:
            r = http_requests.post(
                'https://api.resend.com/emails',
                headers={'Authorization': f'Bearer {RESEND_API_KEY}', 'Content-Type': 'application/json'},
                json={
                    'from': 'Сервис решений <onboarding@resend.dev>',
                    'to': [to],
                    'subject': subject,
                    'html': body,
                },
                timeout=15,
            )
            if r.status_code in (200, 201):
                print(f'[Email] sent to {to} via Resend')
                return True, None
            err = (r.json() or {}).get('message', r.text[:200]) if r.text else 'Unknown error'
            return False, err
        except Exception as e:
            return False, str(e)[:200]

    # Mailgun по HTTPS (нужен домен или sandbox)
    if MAILGUN_API_KEY and MAILGUN_DOMAIN and MAILGUN_FROM:
        try:
            r = http_requests.post(
                f'https://api.mailgun.net/v3/{MAILGUN_DOMAIN}/messages',
                auth=('api', MAILGUN_API_KEY),
                data={
                    'from': formataddr(('Сервис решений', MAILGUN_FROM)),
                    'to': to,
                    'subject': subject,
                    'html': body,
                },
                timeout=15,
            )
            if r.status_code == 200:
                print(f'[Email] sent to {to} via Mailgun')
                return True, None
            err = (r.json() or {}).get('message', r.text[:200]) if r.text else 'Unknown error'
            return False, err
        except Exception as e:
            return False, str(e)[:200]

    # Gmail SMTP — как раньше: один раз порт 465, пароль без пробелов
    if not SMTP_EMAIL or not SMTP_PASSWORD:
        return False, 'Почта не настроена. Добавьте RESEND_API_KEY (resend.com, без домена) или SMTP_EMAIL/SMTP_PASSWORD в .env'
    smtp_pass = (SMTP_PASSWORD or '').replace(' ', '')
    msg = MIMEText(body, 'html', 'utf-8')
    msg['From'] = formataddr(('Сервис решений', SMTP_EMAIL))
    msg['To'] = to
    msg['Subject'] = Header(subject, 'utf-8')
    msg['Date'] = datetime.utcnow().strftime('%a, %d %b %Y %H:%M:%S +0000')
    try:
        import smtplib
        with smtplib.SMTP_SSL('smtp.gmail.com', 465, timeout=25) as srv:
            srv.login(SMTP_EMAIL, smtp_pass)
            srv.sendmail(SMTP_EMAIL, [to], msg.as_string())
        print(f'[Email] sent to {to}')
        return True, None
    except smtplib.SMTPAuthenticationError as e:
        print(f'[Email] SMTP auth error: {e}')
        return False, 'Ошибка входа в почту (проверьте пароль приложения Gmail)'
    except Exception as e:
        err = str(e)
        if 'Authentication' in err or '535' in err:
            return False, 'Неверный логин/пароль приложения Gmail'
        if '10054' in err or 'forcibly closed' in err.lower() or 'удаленный хост' in err.lower():
            return False, 'С твоего интернета Gmail недоступен (порт режут). Добавь в .env: MAILGUN_API_KEY и MAILGUN_DOMAIN — см. EMAIL_SETUP.md (mailgun.com, бесплатно).'
        return False, err[:200] if len(err) > 200 else err

@app.route('/test-email')
def test_email():
    """Проверка отправки почты только с localhost (открой http://localhost:5000/test-email?to=твой@email)"""
    if request.remote_addr not in ('127.0.0.1', '::1'):
        return f'<p>Доступно только с localhost. Открой в браузере: <a href="http://127.0.0.1:5000/test-email?to={request.args.get("to", "")}">http://127.0.0.1:5000/test-email?to=твой@email</a></p>', 403
    to = (request.args.get('to') or '').strip() or SMTP_EMAIL
    if not to:
        return '<p>Укажи адрес: ?to=твой@email</p>', 400
    ok, err = _send_email(to, 'Тест почты (localhost)', '<p>Если ты это видишь — отправка с твоего ПК работает.</p>')
    if ok:
        return f'<p>Письмо отправлено на {to}. Проверь папку «Входящие» и «Спам».</p>'
    return f'<p>Ошибка: {err}</p>', 500

@app.route('/api/forgot-password', methods=['POST'])
def forgot_password():
    d = request.get_json() or {}
    email = (d.get('email') or '').strip().lower()
    if not email:
        return jsonify({"error": "Укажите email"}), 400
    conn = get_db()
    user = conn.execute("SELECT id FROM users WHERE email=?", (email,)).fetchone()
    conn.close()
    if not user:
        return jsonify({"ok": True})
    code = str(secrets.randbelow(900000) + 100000)
    _reset_codes[email] = {'code': code, 'ts': datetime.utcnow(), 'user_id': user['id']}
    print(f'[Password Reset] code for {email}: {code}')
    ok, err_msg = _send_email(email,
        'Код восстановления пароля',
        f'<div style="font-family:sans-serif;max-width:480px;margin:0 auto;padding:2rem;">'
        f'<h2 style="color:#4f46e5;">Восстановление пароля</h2>'
        f'<p>Ваш код подтверждения:</p>'
        f'<div style="font-size:2rem;font-weight:700;letter-spacing:0.3em;background:#f3f4f6;'
        f'padding:1rem;border-radius:0.75rem;text-align:center;margin:1rem 0;">{code}</div>'
        f'<p style="color:#6b7280;font-size:0.875rem;">Код действует 10 минут. '
        f'Если вы не запрашивали восстановление — проигнорируйте это письмо.</p></div>')
    if not ok:
        return jsonify({"error": err_msg or "Не удалось отправить письмо. Проверьте настройки почты на сервере."}), 503
    return jsonify({"ok": True})

@app.route('/api/reset-password', methods=['POST'])
def reset_password():
    d = request.get_json() or {}
    email = (d.get('email') or '').strip().lower()
    code = (d.get('code') or '').strip()
    new_pw = d.get('password') or ''
    if not email or not code or not new_pw:
        return jsonify({"error": "Заполните все поля"}), 400
    if len(new_pw) < 6:
        return jsonify({"error": "Пароль — не менее 6 символов"}), 400
    entry = _reset_codes.get(email)
    if not entry or entry['code'] != code:
        return jsonify({"error": "Неверный код"}), 400
    elapsed = (datetime.utcnow() - entry['ts']).total_seconds()
    if elapsed > 600:
        _reset_codes.pop(email, None)
        return jsonify({"error": "Код истёк, запросите новый"}), 400
    conn = get_db()
    conn.execute("UPDATE users SET password_hash=? WHERE id=?", (hash_password(new_pw), entry['user_id']))
    conn.commit()
    conn.close()
    _reset_codes.pop(email, None)
    return jsonify({"ok": True})

@app.route('/api/delete-account', methods=['POST'])
@login_required
def delete_account():
    d = request.get_json() or {}
    if d.get('confirm') != 'Удалить':
        return jsonify({"error": "Введите слово «Удалить» для подтверждения"}), 400
    uid = session['user_id']
    conn = get_db()
    conn.execute("DELETE FROM users WHERE id=?", (uid,))
    conn.commit()
    conn.close()
    session.clear()
    return jsonify({"ok": True})

def _oauth_login(provider, oauth_id, email, first_name, last_name, birthday=''):
    """Find existing user by OAuth or redirect to register with prefilled data."""
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE oauth_provider=? AND oauth_id=?", (provider, oauth_id)).fetchone()
    if not user and email:
        user = conn.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
        if user:
            conn.execute("UPDATE users SET oauth_provider=?, oauth_id=? WHERE id=?", (provider, oauth_id, user['id']))
            conn.commit()
    conn.close()
    if user:
        session['user_id'] = user['id']
        return redirect('/feed.html')
    session['oauth_pending'] = {'provider': provider, 'oauth_id': oauth_id,
                                'email': email, 'firstName': first_name, 'lastName': last_name}
    p = {'email': email, 'firstName': first_name, 'lastName': last_name, 'oauth': provider}
    if birthday:
        p['birthdate'] = birthday
    return redirect('/register.html?' + urlencode(p))


@app.route('/auth/yandex')
def auth_yandex():
    if not YANDEX_CLIENT_ID:
        return '<p>Яндекс ID не настроен. Укажите YANDEX_CLIENT_ID и YANDEX_CLIENT_SECRET в .env</p>', 501
    params = urlencode({'response_type': 'code', 'client_id': YANDEX_CLIENT_ID,
                        'redirect_uri': _oauth_origin() + '/auth/yandex/callback'})
    return redirect('https://oauth.yandex.ru/authorize?' + params)

@app.route('/auth/yandex/callback')
def auth_yandex_callback():
    code = request.args.get('code')
    if not code:
        return redirect('/')
    try:
        r = http_requests.post('https://oauth.yandex.ru/token', data={
            'grant_type': 'authorization_code', 'code': code,
            'client_id': YANDEX_CLIENT_ID, 'client_secret': YANDEX_CLIENT_SECRET
        }, timeout=10)
        token = r.json().get('access_token')
        info = http_requests.get('https://login.yandex.ru/info?format=json',
                                  headers={'Authorization': 'OAuth ' + token}, timeout=10).json()
        print(f'[Yandex info] birthday={info.get("birthday")!r}  keys={list(info.keys())}')
        return _oauth_login('yandex', info['id'],
                            info.get('default_email', ''),
                            info.get('first_name', ''),
                            info.get('last_name', ''),
                            birthday=info.get('birthday') or '')
    except Exception as e:
        print(f'[Yandex OAuth error] {e}')
        return redirect('/')


@app.route('/auth/vk')
def auth_vk():
    if not VK_CLIENT_ID:
        return '<p>VK ID не настроен. Укажите VK_CLIENT_ID и VK_CLIENT_SECRET в .env</p>', 501
    import hashlib, base64
    code_verifier = secrets.token_urlsafe(64)[:128]
    digest = hashlib.sha256(code_verifier.encode('ascii')).digest()
    code_challenge = base64.urlsafe_b64encode(digest).rstrip(b'=').decode('ascii')
    session['vk_code_verifier'] = code_verifier
    state = secrets.token_urlsafe(32)
    session['vk_state'] = state
    params = urlencode({
        'response_type': 'code',
        'client_id': VK_CLIENT_ID,
        'redirect_uri': _oauth_origin() + '/auth/vk/callback',
        'code_challenge': code_challenge,
        'code_challenge_method': 'S256',
        'state': state,
        'scope': '',
    })
    return redirect('https://id.vk.com/authorize?' + params)

@app.route('/auth/vk/callback')
def auth_vk_callback():
    code = request.args.get('code', '')
    device_id = request.args.get('device_id', '')
    state = request.args.get('state', '')
    if not code:
        return redirect('/')
    saved_state = session.pop('vk_state', '')
    code_verifier = session.pop('vk_code_verifier', '')
    if state != saved_state:
        print(f'[VK OAuth] state mismatch: {state!r} != {saved_state!r}')
        return redirect('/')
    try:
        r = http_requests.post('https://id.vk.com/oauth2/auth', data={
            'grant_type': 'authorization_code',
            'code': code,
            'client_id': VK_CLIENT_ID,
            'code_verifier': code_verifier,
            'device_id': device_id,
            'redirect_uri': _oauth_origin() + '/auth/vk/callback',
            'state': state,
        }, timeout=10).json()
        print(f'[VK token response] {r}')
        token = r.get('access_token', '')
        vk_uid = str(r.get('user_id', ''))
        email = r.get('email', '')
        first_name = ''
        last_name = ''
        if token:
            info = http_requests.post('https://id.vk.com/oauth2/user_info', data={
                'access_token': token, 'client_id': VK_CLIENT_ID,
            }, timeout=10).json()
            print(f'[VK user_info] {info}')
            user = info.get('user', {})
            first_name = user.get('first_name', '')
            last_name = user.get('last_name', '')
            email = email or user.get('email', '')
            if not vk_uid:
                vk_uid = str(user.get('user_id', ''))
            birthday = user.get('birthday', '')
            if birthday and '.' in birthday:
                parts = birthday.split('.')
                if len(parts) == 3:
                    birthday = f'{parts[2]}-{parts[1]}-{parts[0]}'
            return _oauth_login('vk', vk_uid, email, first_name, last_name, birthday=birthday)
        return _oauth_login('vk', vk_uid, email, first_name, last_name)
    except Exception as e:
        print(f'[VK OAuth error] {e}')
        return redirect('/')


@app.route('/auth/google')
def auth_google():
    if not GOOGLE_CLIENT_ID:
        return '<p>Google OAuth не настроен. Укажите GOOGLE_CLIENT_ID и GOOGLE_CLIENT_SECRET в .env</p>', 501
    params = urlencode({'client_id': GOOGLE_CLIENT_ID, 'response_type': 'code',
                        'redirect_uri': _oauth_origin() + '/auth/google/callback',
                        'scope': 'openid email profile', 'access_type': 'offline'})
    return redirect('https://accounts.google.com/o/oauth2/v2/auth?' + params)

@app.route('/auth/google/callback')
def auth_google_callback():
    code = request.args.get('code')
    if not code:
        return redirect('/')
    try:
        r = http_requests.post('https://oauth2.googleapis.com/token', data={
            'code': code, 'client_id': GOOGLE_CLIENT_ID,
            'client_secret': GOOGLE_CLIENT_SECRET,
            'redirect_uri': _oauth_origin() + '/auth/google/callback',
            'grant_type': 'authorization_code'
        }, timeout=10).json()
        token = r.get('access_token')
        info = http_requests.get('https://www.googleapis.com/oauth2/v2/userinfo',
                                  headers={'Authorization': 'Bearer ' + token}, timeout=10).json()
        return _oauth_login('google', info['id'],
                            info.get('email', ''),
                            info.get('given_name', ''),
                            info.get('family_name', ''))
    except Exception as e:
        print(f'[Google OAuth error] {e}')
        return redirect('/')


@app.route('/auth/telegram')
def auth_telegram():
    if not TELEGRAM_BOT_TOKEN:
        return '<p>Telegram не настроен. Укажите TELEGRAM_BOT_TOKEN в .env</p>', 501
    bot_id = TELEGRAM_BOT_TOKEN.split(':')[0]
    origin = request.host_url.rstrip('/')
    redirect_url = origin + '/auth/telegram/callback'
    params = urlencode({'bot_id': bot_id, 'origin': origin,
                        'embed': '0', 'request_access': 'write',
                        'return_to': redirect_url})
    return redirect('https://oauth.telegram.org/auth?' + params)

@app.route('/auth/telegram/callback')
def auth_telegram_callback():
    tg_data = {}
    for key in ('id', 'first_name', 'last_name', 'username', 'photo_url', 'auth_date', 'hash'):
        val = request.args.get(key, '')
        if val:
            tg_data[key] = val
    tg_hash = tg_data.pop('hash', '')
    if not tg_hash or not tg_data.get('id'):
        return redirect('/')
    check_string = '\n'.join(f'{k}={tg_data[k]}' for k in sorted(tg_data.keys()))
    secret_key = hashlib.sha256(TELEGRAM_BOT_TOKEN.encode('utf-8')).digest()
    expected_hash = hmac.HMAC(secret_key, check_string.encode('utf-8'), hashlib.sha256).hexdigest()
    if expected_hash != tg_hash:
        print(f'[Telegram] hash mismatch')
        return redirect('/')
    return _oauth_login('telegram', tg_data['id'], '',
                        tg_data.get('first_name', ''), tg_data.get('last_name', ''))


@app.route('/oauth-setup')
def oauth_setup():
    """Страница с актуальными URI для Google и Яндекс — копируй и добавляй в настройки OAuth."""
    base = _oauth_origin()
    yandex_callback = base + '/auth/yandex/callback'
    google_origin = base
    google_callback = base + '/auth/google/callback'
    html = '''<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>OAuth URI для Google и Яндекс</title>
<style>body{font-family:sans-serif;max-width:640px;margin:2rem auto;padding:0 1rem;}
h1{font-size:1.25rem;} .box{background:#f5f5f5;padding:1rem;border-radius:8px;margin:1rem 0;}
code{background:#eee;padding:2px 6px;border-radius:4px;word-break:break-all;}
.copy{display:block;margin-top:0.5rem;cursor:pointer;color:#06c;} .warn{color:#c00;margin-top:1rem;}
.note{background:#fff3cd;padding:0.75rem;border-radius:6px;margin:1rem 0;font-size:0.9rem;}
</style></head><body>
<h1>URI для Google и Яндекс</h1>
<p>Скопируй адреса ниже и добавь их в настройках OAuth. Чтобы при каждом перезапуске туннеля не менять настройки — задай в .env <code>OAUTH_BASE_URL</code> и один раз добавь эти URI.</p>
<p><strong>Базовый URL → в .env как OAUTH_BASE_URL (без слэша в конце), затем перезапусти Flask:</strong></p>
<div class="box"><code id="base">''' + base + '''</code> <small>(клик — копировать)</small></div>
<h2>Яндекс (oauth.yandex.ru → Платформы → Веб-сервисы)</h2>
<p>Suggest Hostname (хост страницы с кнопкой):</p>
<div class="box"><code id="ya-host">''' + base + '''</code></div>
<p>Redirect URI (Callback):</p>
<div class="box"><code id="ya">''' + yandex_callback + '''</code></div>
<h2>Google (APIs & Services → Credentials → OAuth 2.0 client)</h2>
<p>Авторизованные источники JavaScript:</p>
<div class="box"><code id="go-origin">''' + google_origin + '''</code></div>
<p>Авторизованные URI перенаправления:</p>
<div class="box"><code id="go-cb">''' + google_callback + '''</code></div>
<div class="note">После смены URI в Google/Яндексе изменения могут применяться до нескольких часов. Старые URI (localhost и прошлые туннели) можно не удалять — тогда и локально, и по туннелю вход будет работать.</div>
<p class="warn">В .env пропиши: OAUTH_BASE_URL=''' + base + ''' и перезапусти Flask. Тогда эти URI не будут «уплывать» при следующем запросе.</p>
<script>
function copy(id){navigator.clipboard.writeText(document.getElementById(id).textContent);alert("Скопировано");}
["base","ya-host","ya","go-origin","go-cb"].forEach(function(id){var el=document.getElementById(id);if(el)el.onclick=function(){copy(id);};});
</script>
</body></html>'''
    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}


@app.route('/api/me')
@login_required
def me():
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE id=?", (session['user_id'],)).fetchone()
    conn.close()
    if not user: session.clear(); return jsonify({"error": "Пользователь не найден"}), 401
    return jsonify({"user": user_dict(user)})


# ═══════════════════════════════════════════════════════════════
#  POLLS
# ═══════════════════════════════════════════════════════════════
def _ensure_is_public_columns(conn):
    """Добавить колонку is_public, если её нет (для старых БД)."""
    try:
        conn.execute("SELECT is_public FROM polls LIMIT 1")
    except sqlite3.OperationalError:
        try:
            conn.execute("ALTER TABLE polls ADD COLUMN is_public INTEGER DEFAULT 0")
            conn.execute("ALTER TABLE decisions ADD COLUMN is_public INTEGER DEFAULT 0")
            conn.commit()
        except sqlite3.OperationalError:
            pass


@app.route('/api/polls', methods=['POST'])
@block_if_banned
def create_poll():
    d = request.get_json() or {}
    title = (d.get('title') or '').strip()
    desc = (d.get('description') or '').strip()
    options = d.get('options') or []
    multiple = 1 if d.get('multipleChoice') else 0
    auth_only = 1 if d.get('authOnly') else 0
    show_results = d.get('showResults') or 'always'
    show_voters = 1 if d.get('showVoters') else 0
    anonymous = 1 if d.get('anonymous') else 0
    if anonymous: show_voters = 0
    elif show_voters: anonymous = 0
    max_votes = int(d.get('maxVotes') or 0)
    is_public = 1 if (d.get('isPublic') is True or d.get('isPublic') == 'true') else 0
    deadline = (d.get('deadline') or '').strip() or None
    if not title: return jsonify({"error": "Введите тему"}), 400
    if len(options) < 2: return jsonify({"error": "Минимум 2 варианта"}), 400
    mod_reason = None
    ok, reason = _moderate_content(title, desc, options)
    if not ok:
        w, blocked = _apply_warning(session['user_id'], 'poll', title, desc, reason)
        if blocked:
            return jsonify({"error": "Аккаунт заблокирован за повторные нарушения", "blocked": True}), 403
        return jsonify({"error": reason or "Контент отклонён модерацией", "warning": True, "reason": reason, "warningCount": w}), 400
    slug = uuid.uuid4().hex[:10]
    conn = get_db()
    _ensure_is_public_columns(conn)
    cur = conn.execute("INSERT INTO polls (user_id,title,description,slug,multiple_choice,auth_only,show_results,deadline,show_voters,anonymous,max_votes,is_public) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                       (session['user_id'], title, desc, slug, multiple, auth_only, show_results, deadline, show_voters, anonymous, max_votes, is_public))
    pid = cur.lastrowid
    for i, opt in enumerate(options):
        t = (opt if isinstance(opt, str) else opt.get('text', '')).strip()
        if t: conn.execute("INSERT INTO poll_options (poll_id,text,sort_order) VALUES (?,?,?)", (pid, t, i))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "slug": slug, "isPublic": bool(is_public)}), 201

@app.route('/api/polls/<slug>', methods=['PUT'])
@block_if_banned
def edit_poll(slug):
    d = request.get_json() or {}
    conn = get_db()
    _ensure_is_public_columns(conn)
    p = conn.execute("SELECT * FROM polls WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    title = (d.get('title') or p['title']).strip()
    desc = d.get('description') if d.get('description') is not None else p['description']
    deadline = d.get('deadline') if d.get('deadline') is not None else p['deadline']
    show_results = d.get('showResults') or p['show_results']
    show_voters = int(d.get('showVoters', p['show_voters']))
    anonymous = int(d.get('anonymous', p['anonymous']))
    if anonymous: show_voters = 0
    elif show_voters: anonymous = 0
    max_votes = int(d.get('maxVotes') or p['max_votes'] or 0)
    is_public = 1 if (d.get('isPublic') is True or d.get('isPublic') == 'true') else (0 if 'isPublic' in d else (int(p['is_public']) if 'is_public' in p.keys() else 0))
    edit_opts = d.get('options') or []
    ok, reason = _moderate_content(title, desc, edit_opts)
    if not ok:
        w, blocked = _apply_warning(session['user_id'], 'poll', title, desc, reason)
        if blocked:
            conn.close()
            return jsonify({"error": "Аккаунт заблокирован за повторные нарушения", "blocked": True}), 403
        conn.close()
        return jsonify({"error": reason or "Контент отклонён модерацией", "warning": True, "reason": reason, "warningCount": w}), 400
    conn.execute("UPDATE polls SET title=?,description=?,deadline=?,show_results=?,show_voters=?,anonymous=?,max_votes=?,is_public=? WHERE id=?",
                 (title, desc, deadline or None, show_results, show_voters, anonymous, max_votes, is_public, p['id']))
    if 'options' in d:
        conn.execute("DELETE FROM poll_options WHERE poll_id=?", (p['id'],))
        conn.execute("DELETE FROM poll_votes WHERE poll_id=? AND option_id NOT IN (SELECT id FROM poll_options WHERE poll_id=?)", (p['id'], p['id']))
        for i, opt in enumerate(d['options']):
            t = (opt if isinstance(opt, str) else opt.get('text', '')).strip()
            if t: conn.execute("INSERT INTO poll_options (poll_id,text,sort_order) VALUES (?,?,?)", (p['id'], t, i))
    conn.commit(); conn.close()
    return jsonify({"ok": True})

@app.route('/api/polls')
@login_required
def my_polls():
    conn = get_db()
    polls = conn.execute("SELECT * FROM polls WHERE user_id=? ORDER BY created_at DESC", (session['user_id'],)).fetchall()
    result = []
    for p in polls:
        opts = conn.execute("SELECT * FROM poll_options WHERE poll_id=? ORDER BY sort_order", (p['id'],)).fetchall()
        tv = conn.execute("SELECT COUNT(DISTINCT fingerprint) as c FROM poll_votes WHERE poll_id=?", (p['id'],)).fetchone()['c']
        active = bool(p['is_active']) and check_deadline(p['deadline'])
        result.append({"id": p['id'], "title": p['title'], "description": p['description'],
            "slug": p['slug'], "multipleChoice": bool(p['multiple_choice']),
            "isActive": active, "authOnly": bool(p['auth_only']),
            "showResults": p['show_results'], "deadline": p['deadline'],
            "showVoters": bool(p['show_voters']), "anonymous": bool(p['anonymous']),
            "maxVotes": p['max_votes'] or 0,
            "isPublic": bool(p['is_public'] if 'is_public' in p.keys() else 0),
            "createdAt": p['created_at'], "totalVotes": tv,
            "options": [{"id": o['id'], "text": o['text']} for o in opts]})
    conn.close()
    return jsonify({"polls": result})

@app.route('/api/polls/<slug>')
def get_poll(slug):
    conn = get_db()
    p = conn.execute("SELECT * FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    author = conn.execute("SELECT * FROM users WHERE id=?", (p['user_id'],)).fetchone()
    opts = conn.execute("SELECT * FROM poll_options WHERE poll_id=? ORDER BY sort_order", (p['id'],)).fetchall()
    active = bool(p['is_active']) and check_deadline(p['deadline'])
    uid = session.get('user_id')
    has_voted = False
    if uid:
        has_voted = conn.execute("SELECT 1 FROM poll_votes WHERE poll_id=? AND user_id=?", (p['id'], uid)).fetchone() is not None
    conn.close()
    return jsonify({"poll": {"id": p['id'], "title": p['title'], "description": p['description'],
        "slug": p['slug'], "multipleChoice": bool(p['multiple_choice']), "isActive": active,
        "authOnly": bool(p['auth_only']), "showResults": p['show_results'], "deadline": p['deadline'],
        "showVoters": bool(p['show_voters']), "anonymous": bool(p['anonymous']),
        "maxVotes": p['max_votes'] or 0, "hasVoted": has_voted,
        "author": {"firstName": author['first_name'], "lastName": author['last_name']},
        "options": [{"id": o['id'], "text": o['text']} for o in opts]}})

@app.route('/api/polls/<slug>/vote', methods=['POST'])
def vote_poll(slug):
    uid = session.get('user_id')
    if uid:
        conn_chk = get_db()
        u_chk = conn_chk.execute("SELECT is_blocked FROM users WHERE id=?", (uid,)).fetchone()
        conn_chk.close()
        if u_chk and u_chk['is_blocked']:
            return jsonify({"error": "Аккаунт заблокирован. Голосование недоступно.", "blocked": True}), 403
    d = request.get_json()
    oids = d.get('optionIds') or []
    if not oids: return jsonify({"error": "Выберите вариант"}), 400
    conn = get_db()
    p = conn.execute("SELECT * FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    if not (bool(p['is_active']) and check_deadline(p['deadline'])):
        conn.close(); return jsonify({"error": "Голосование завершено"}), 400
    if p['auth_only'] and 'user_id' not in session:
        conn.close(); return jsonify({"error": "Голосовать могут только авторизованные пользователи"}), 401
    fp = get_fingerprint()
    uid = session.get('user_id')
    ex = None
    if uid:
        ex = conn.execute("SELECT id FROM poll_votes WHERE poll_id=? AND user_id=?", (p['id'], uid)).fetchone()
    if not ex:
        ex = conn.execute("SELECT id FROM poll_votes WHERE poll_id=? AND fingerprint=?", (p['id'], fp)).fetchone()
    if ex: conn.close(); return jsonify({"error": "Вы уже голосовали"}), 409
    max_v = p['max_votes'] or 0
    if max_v > 0:
        cur_total = conn.execute("SELECT COUNT(DISTINCT COALESCE(user_id,fingerprint)) as c FROM poll_votes WHERE poll_id=?", (p['id'],)).fetchone()['c']
        if cur_total >= max_v:
            conn.close(); return jsonify({"error": "Достигнут лимит голосов"}), 400
    for oid in oids:
        conn.execute("INSERT INTO poll_votes (poll_id,option_id,user_id,fingerprint) VALUES (?,?,?,?)", (p['id'], oid, uid, fp))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "fingerprint": fp})

@app.route('/api/polls/<slug>/results')
def poll_results(slug):
    conn = get_db()
    p = conn.execute("SELECT * FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    sr = p['show_results']
    active = bool(p['is_active']) and check_deadline(p['deadline'])
    is_author = session.get('user_id') == p['user_id']
    if sr == 'author_only' and not is_author:
        conn.close(); return jsonify({"error": "Результаты доступны только автору"}), 403
    if sr == 'after' and active and not is_author:
        conn.close(); return jsonify({"error": "Результаты будут доступны после завершения"}), 403
    opts = conn.execute("SELECT * FROM poll_options WHERE poll_id=? ORDER BY sort_order", (p['id'],)).fetchall()
    total = conn.execute("SELECT COUNT(DISTINCT COALESCE(user_id,fingerprint)) as c FROM poll_votes WHERE poll_id=?", (p['id'],)).fetchone()['c']
    results = []
    for o in opts:
        cnt = conn.execute("SELECT COUNT(*) as c FROM poll_votes WHERE option_id=?", (o['id'],)).fetchone()['c']
        voters = []
        if is_author and p['show_voters'] and not p['anonymous']:
            vrows = conn.execute("""SELECT u.first_name, u.last_name FROM poll_votes pv
                JOIN users u ON pv.user_id = u.id WHERE pv.option_id=?""", (o['id'],)).fetchall()
            voters = [r['first_name'] + ' ' + r['last_name'] for r in vrows]
        results.append({"id": o['id'], "text": o['text'], "votes": cnt, "voters": voters})
    conn.close()
    return jsonify({"title": p['title'], "totalVoters": total, "results": results,
        "showResults": sr, "isActive": active,
        "showVoters": bool(p['show_voters']) and is_author, "anonymous": bool(p['anonymous'])})

@app.route('/api/polls/<slug>/toggle', methods=['POST'])
@block_if_banned
def toggle_poll(slug):
    conn = get_db()
    p = conn.execute("SELECT * FROM polls WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    ns = 0 if p['is_active'] else 1
    conn.execute("UPDATE polls SET is_active=? WHERE id=?", (ns, p['id']))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "isActive": bool(ns)})

@app.route('/api/polls/<slug>/publish-to-feed', methods=['POST'])
@block_if_banned
def poll_publish_to_feed(slug):
    conn = get_db()
    _ensure_is_public_columns(conn)
    p = conn.execute("SELECT * FROM polls WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("UPDATE polls SET is_public=1 WHERE id=?", (p['id'],))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "isPublic": True})

@app.route('/api/polls/<slug>/unpublish-from-feed', methods=['POST'])
@block_if_banned
def poll_unpublish_from_feed(slug):
    conn = get_db()
    p = conn.execute("SELECT * FROM polls WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("UPDATE polls SET is_public=0 WHERE id=?", (p['id'],))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "isPublic": False})

@app.route('/api/polls/<slug>/qr')
def poll_qr(slug):
    url = request.host_url.rstrip('/') + '/vote.html?id=' + slug
    img = qrcode.make(url, box_size=8, border=2)
    buf = io.BytesIO(); img.save(buf, format='PNG')
    return jsonify({"qr": "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode(), "url": url})


# ═══════════════════════════════════════════════════════════════
#  DECISIONS
# ═══════════════════════════════════════════════════════════════
@app.route('/api/decisions', methods=['POST'])
@block_if_banned
def create_decision():
    d = request.get_json() or {}
    title = (d.get('title') or '').strip()
    desc = (d.get('description') or '').strip()
    alts = d.get('alternatives') or []
    crits = d.get('criteria') or []
    auth_only = 1 if d.get('authOnly') else 0
    show_respondents = 1 if d.get('showRespondents') else 0
    anonymous = 1 if d.get('anonymous') else 0
    if anonymous: show_respondents = 0
    elif show_respondents: anonymous = 0
    show_results = d.get('showResults') or 'always'
    deadline = (d.get('deadline') or '').strip() or None
    is_public = 1 if (d.get('isPublic') is True or d.get('isPublic') == 'true') else 0
    scale_max = int(d.get('scaleMax') or d.get('scale_max') or 5)
    if scale_max not in (3, 5, 10, 100): scale_max = 5
    if not title: return jsonify({"error": "Введите вопрос"}), 400
    if len(alts) < 2: return jsonify({"error": "Минимум 2 варианта"}), 400
    if len(crits) < 1: return jsonify({"error": "Минимум 1 критерий"}), 400
    mod_reason = None
    mod_warning_count = 0
    if is_public:
        ok, reason = _moderate_content(title, desc, alts + crits)
        if not ok:
            w, blocked = _apply_warning(session['user_id'], 'decision', title, desc, reason)
            if blocked:
                return jsonify({"error": "Аккаунт заблокирован за повторные нарушения", "blocked": True}), 403
            is_public = 0
            mod_reason = reason
            mod_warning_count = w
    slug = uuid.uuid4().hex[:10]
    conn = get_db()
    _ensure_is_public_columns(conn)
    cur = conn.execute("INSERT INTO decisions (user_id,title,description,slug,auth_only,deadline,show_respondents,anonymous,show_results,is_public,scale_max) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                       (session['user_id'], title, desc, slug, auth_only, deadline, show_respondents, anonymous, show_results, is_public, scale_max))
    did = cur.lastrowid
    for i, a in enumerate(alts):
        n = (a if isinstance(a, str) else a.get('name', '')).strip()
        if n: conn.execute("INSERT INTO decision_alternatives (decision_id,name,sort_order) VALUES (?,?,?)", (did, n, i))
    for i, c in enumerate(crits):
        n = (c if isinstance(c, str) else c.get('name', '')).strip()
        if n: conn.execute("INSERT INTO decision_criteria (decision_id,name,sort_order) VALUES (?,?,?)", (did, n, i))
    conn.commit(); conn.close()
    out = {"ok": True, "slug": slug, "isPublic": bool(is_public)}
    if mod_reason:
        out["reason"] = mod_reason
        out["notPublishedToFeed"] = True
        out["warning"] = True
        out["warningCount"] = mod_warning_count
    return jsonify(out), 201

@app.route('/api/decisions')
@login_required
def my_decisions():
    conn = get_db()
    decs = conn.execute("SELECT * FROM decisions WHERE user_id=? ORDER BY created_at DESC", (session['user_id'],)).fetchall()
    result = []
    for dc in decs:
        ac = conn.execute("SELECT COUNT(*) as c FROM decision_alternatives WHERE decision_id=?", (dc['id'],)).fetchone()['c']
        cc = conn.execute("SELECT COUNT(*) as c FROM decision_criteria WHERE decision_id=?", (dc['id'],)).fetchone()['c']
        rc = conn.execute("SELECT COUNT(*) as c FROM decision_responses WHERE decision_id=?", (dc['id'],)).fetchone()['c']
        active = bool(dc['is_active']) and check_deadline(dc['deadline'])
        alts = conn.execute("SELECT name FROM decision_alternatives WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
        crits_list = conn.execute("SELECT name FROM decision_criteria WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
        scale_max = int(dc['scale_max']) if 'scale_max' in dc.keys() and dc['scale_max'] else 5
        result.append({"id": dc['id'], "title": dc['title'], "description": dc['description'] or '',
            "slug": dc['slug'], "createdAt": dc['created_at'], "alternativesCount": ac, "criteriaCount": cc,
            "responsesCount": rc, "isActive": active, "deadline": dc['deadline'],
            "authOnly": bool(dc['auth_only']),
            "showRespondents": bool(dc['show_respondents'] if 'show_respondents' in dc.keys() else 0),
            "anonymous": bool(dc['anonymous'] if 'anonymous' in dc.keys() else 0),
            "showResults": (dc['show_results'] if 'show_results' in dc.keys() else 'always') or 'always',
            "isPublic": bool(dc['is_public'] if 'is_public' in dc.keys() else 0),
            "scaleMax": scale_max,
            "alternatives": [a['name'] for a in alts],
            "criteria": [c['name'] for c in crits_list]})
    conn.close()
    return jsonify({"decisions": result})

@app.route('/api/decisions/<slug>')
def get_decision(slug):
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    author = conn.execute("SELECT * FROM users WHERE id=?", (dc['user_id'],)).fetchone()
    alts = conn.execute("SELECT * FROM decision_alternatives WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
    crits = conn.execute("SELECT * FROM decision_criteria WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
    active = bool(dc['is_active']) and check_deadline(dc['deadline'])
    uid = session.get('user_id')
    has_responded = False
    if uid:
        has_responded = conn.execute("SELECT 1 FROM decision_responses WHERE decision_id=? AND user_id=?", (dc['id'], uid)).fetchone() is not None
    conn.close()
    sr = (dc['show_results'] if 'show_results' in dc.keys() else 'always') or 'always'
    return jsonify({"decision": {"id": dc['id'], "title": dc['title'], "description": dc['description'],
        "slug": dc['slug'], "isActive": active, "authOnly": bool(dc['auth_only']), "deadline": dc['deadline'],
        "showResults": sr, "hasResponded": has_responded,
        "author": {"firstName": author['first_name'], "lastName": author['last_name']},
        "alternatives": [{"id": a['id'], "name": a['name']} for a in alts],
        "criteria": [{"id": c['id'], "name": c['name']} for c in crits],
        "scaleMax": int(dc['scale_max']) if 'scale_max' in dc.keys() and dc['scale_max'] else 5}})

@app.route('/api/decisions/<slug>/respond', methods=['POST'])
def respond_decision(slug):
    uid = session.get('user_id')
    if uid:
        conn_chk = get_db()
        u_chk = conn_chk.execute("SELECT is_blocked FROM users WHERE id=?", (uid,)).fetchone()
        conn_chk.close()
        if u_chk and u_chk['is_blocked']:
            return jsonify({"error": "Аккаунт заблокирован. Оценка недоступна.", "blocked": True}), 403
    d = request.get_json()
    scores = d.get('scores') or {}
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    if not (bool(dc['is_active']) and check_deadline(dc['deadline'])):
        conn.close(); return jsonify({"error": "Сбор ответов завершён"}), 400
    if dc['auth_only'] and 'user_id' not in session:
        conn.close(); return jsonify({"error": "Только для авторизованных"}), 401
    fp = get_fingerprint()
    uid = session.get('user_id')
    if uid:
        ex = conn.execute("SELECT id FROM decision_responses WHERE decision_id=? AND user_id=?", (dc['id'], uid)).fetchone()
    else:
        ex = conn.execute("SELECT id FROM decision_responses WHERE decision_id=? AND fingerprint=?", (dc['id'], fp)).fetchone()
    if ex: conn.close(); return jsonify({"error": "Вы уже ответили"}), 409
    cur = conn.execute("INSERT INTO decision_responses (decision_id,user_id,fingerprint) VALUES (?,?,?)", (dc['id'], uid, fp))
    rid = cur.lastrowid
    scale_max = int(dc['scale_max']) if 'scale_max' in dc.keys() and dc['scale_max'] else 5
    for key, val in scores.items():
        parts = key.split('_')
        if len(parts) == 2:
            v = int(val)
            if v < 1 or v > scale_max: v = max(1, min(scale_max, v))
            conn.execute("INSERT INTO decision_scores (response_id,alternative_id,criterion_id,score) VALUES (?,?,?,?)",
                         (rid, int(parts[0]), int(parts[1]), v))
    conn.commit(); conn.close()
    return jsonify({"ok": True})

@app.route('/api/decisions/<slug>/results')
def decision_results(slug):
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    sr = (dc['show_results'] if 'show_results' in dc.keys() else 'always') or 'always'
    active = bool(dc['is_active']) and check_deadline(dc['deadline'])
    is_author = session.get('user_id') == dc['user_id']
    if sr == 'author_only' and not is_author:
        conn.close(); return jsonify({"error": "Результаты доступны только автору"}), 403
    if sr == 'after' and active and not is_author:
        conn.close(); return jsonify({"error": "Результаты будут доступны после завершения"}), 403
    alts = conn.execute("SELECT * FROM decision_alternatives WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
    crits = conn.execute("SELECT * FROM decision_criteria WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
    rc = conn.execute("SELECT COUNT(*) as c FROM decision_responses WHERE decision_id=?", (dc['id'],)).fetchone()['c']
    ranking = []
    details = {}
    for a in alts:
        total = 0
        for c in crits:
            avg_row = conn.execute(
                "SELECT AVG(score) as avg FROM decision_scores WHERE alternative_id=? AND criterion_id=?",
                (a['id'], c['id'])).fetchone()
            avg = round(avg_row['avg'], 2) if avg_row['avg'] else 0
            details[str(a['id']) + '_' + str(c['id'])] = avg
            total += avg
        avg_total = round(total / len(crits), 2) if crits else 0
        ranking.append({"id": a['id'], "name": a['name'], "score": avg_total})
    ranking.sort(key=lambda x: x['score'], reverse=True)
    respondents = []
    if is_author and dc['show_respondents'] and not dc['anonymous']:
        rrows = conn.execute("""SELECT u.first_name, u.last_name FROM decision_responses dr
            JOIN users u ON dr.user_id = u.id WHERE dr.decision_id=?""", (dc['id'],)).fetchall()
        respondents = [r['first_name'] + ' ' + r['last_name'] for r in rrows]
    conn.close()
    scale_max = int(dc['scale_max']) if 'scale_max' in dc.keys() and dc['scale_max'] else 5
    return jsonify({"title": dc['title'], "totalResponses": rc, "ranking": ranking, "details": details,
        "alternatives": [{"id": a['id'], "name": a['name']} for a in alts],
        "criteria": [{"id": c['id'], "name": c['name']} for c in crits],
        "showRespondents": bool(dc['show_respondents']) and is_author, "anonymous": bool(dc['anonymous']),
        "respondents": respondents, "isAuthor": is_author, "scaleMax": scale_max})

@app.route('/api/decisions/<slug>/qr')
def decision_qr(slug):
    url = request.host_url.rstrip('/') + '/respond.html?id=' + slug
    img = qrcode.make(url, box_size=8, border=2)
    buf = io.BytesIO(); img.save(buf, format='PNG')
    return jsonify({"qr": "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode(), "url": url})

@app.route('/api/decisions/<slug>/toggle', methods=['POST'])
@block_if_banned
def toggle_decision(slug):
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    ns = 0 if dc['is_active'] else 1
    conn.execute("UPDATE decisions SET is_active=? WHERE id=?", (ns, dc['id']))
    conn.commit(); conn.close()
    return jsonify({"ok": True})

@app.route('/api/decisions/<slug>/publish-to-feed', methods=['POST'])
@block_if_banned
def decision_publish_to_feed(slug):
    """Опубликовать решение в ленту (is_public=1)."""
    conn = get_db()
    _ensure_is_public_columns(conn)
    dc = conn.execute("SELECT * FROM decisions WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("UPDATE decisions SET is_public=1 WHERE id=?", (dc['id'],))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "isPublic": True})

@app.route('/api/decisions/<slug>/unpublish-from-feed', methods=['POST'])
@block_if_banned
def decision_unpublish_from_feed(slug):
    """Убрать решение из ленты (is_public=0)."""
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("UPDATE decisions SET is_public=0 WHERE id=?", (dc['id'],))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "isPublic": False})

@app.route('/api/decisions/<slug>', methods=['PUT'])
@block_if_banned
def edit_decision(slug):
    d = request.get_json() or {}
    conn = get_db()
    _ensure_is_public_columns(conn)
    dc = conn.execute("SELECT * FROM decisions WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    title = (d.get('title') or dc['title']).strip()
    desc = d.get('description') if d.get('description') is not None else dc['description']
    auth_only = int(d.get('authOnly', dc['auth_only']))
    show_respondents = int(d.get('showRespondents', dc['show_respondents'] if 'show_respondents' in dc.keys() else 0))
    anonymous = int(d.get('anonymous', dc['anonymous'] if 'anonymous' in dc.keys() else 0))
    if anonymous: show_respondents = 0
    elif show_respondents: anonymous = 0
    show_results = d.get('showResults') or dc['show_results'] if 'show_results' in dc.keys() else 'always'
    deadline = d.get('deadline') if d.get('deadline') is not None else dc['deadline']
    is_public = 1 if (d.get('isPublic') is True or d.get('isPublic') == 'true') else (0 if 'isPublic' in d else (int(dc['is_public']) if 'is_public' in dc.keys() else 0))
    scale_max = int(d.get('scaleMax') or d.get('scale_max') or dc.get('scale_max') or 5)
    if scale_max not in (3, 5, 10, 100): scale_max = 5
    mod_reason = None
    mod_warning_count = 0
    if is_public:
        alts_text = d.get('alternatives') or []
        crits_text = d.get('criteria') or []
        ok, reason = _moderate_content(title, desc, alts_text + crits_text)
        if not ok:
            w, blocked = _apply_warning(session['user_id'], 'decision', title, desc, reason)
            if blocked:
                conn.close()
                return jsonify({"error": "Аккаунт заблокирован за повторные нарушения", "blocked": True}), 403
            is_public = 0
            mod_reason = reason
            mod_warning_count = w
    conn.execute(
        "UPDATE decisions SET title=?,description=?,auth_only=?,deadline=?,show_respondents=?,anonymous=?,show_results=?,is_public=?,scale_max=? WHERE id=?",
        (title, desc, auth_only, deadline or None, show_respondents, anonymous, show_results, is_public, scale_max, dc['id']))
    if 'alternatives' in d:
        conn.execute("DELETE FROM decision_alternatives WHERE decision_id=?", (dc['id'],))
        for i, a in enumerate(d['alternatives']):
            n = (a if isinstance(a, str) else a.get('name', '')).strip()
            if n: conn.execute("INSERT INTO decision_alternatives (decision_id,name,sort_order) VALUES (?,?,?)", (dc['id'], n, i))
    if 'criteria' in d:
        conn.execute("DELETE FROM decision_criteria WHERE decision_id=?", (dc['id'],))
        for i, c in enumerate(d['criteria']):
            n = (c if isinstance(c, str) else c.get('name', '')).strip()
            if n: conn.execute("INSERT INTO decision_criteria (decision_id,name,sort_order) VALUES (?,?,?)", (dc['id'], n, i))
    conn.commit(); conn.close()
    out = {"ok": True}
    if mod_reason:
        out["reason"] = mod_reason
        out["notPublishedToFeed"] = True
        out["warning"] = True
        out["warningCount"] = mod_warning_count
    return jsonify(out)

@app.route('/api/decisions/<slug>', methods=['DELETE'])
@block_if_banned
def delete_decision(slug):
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=? AND user_id=?", (slug, session['user_id'])).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("DELETE FROM decisions WHERE id=?", (dc['id'],))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


# ═══════════════════════════════════════════════════════════════
#  AI
# ═══════════════════════════════════════════════════════════════
_CRITERIA_FALLBACK = ["Стоимость", "Качество", "Удобство", "Надёжность", "Доступность"]

@app.route('/api/ai/suggest-criteria', methods=['POST'])
def suggest_criteria():
    d = request.get_json()
    question = (d.get('question') or '').strip()
    if not question:
        return jsonify({"error": "Введите вопрос"}), 400

    sys_prompt = (
        "Ты помощник для принятия решений. Пользователь задаёт вопрос — предложи 4-6 "
        "критериев оценки, которые реально помогут сравнить варианты именно по этому вопросу. "
        "Критерии должны быть конкретные, уместные и понятные. "
        "Отвечай ТОЛЬКО JSON-массивом строк, без пояснений. "
        'Пример: ["Стоимость", "Удобство", "Качество"]'
    )
    text, source, _ = _call_ai(sys_prompt, question, temperature=0.7, max_tokens=200)
    if text:
        try:
            if text.startswith("```"):
                text = text.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            criteria = json.loads(text)
            if isinstance(criteria, list) and len(criteria) >= 2:
                return jsonify({"criteria": criteria, "source": source})
        except (json.JSONDecodeError, ValueError):
            pass

    return jsonify({"criteria": _CRITERIA_FALLBACK, "source": "local"})


# ═══════════════════════════════════════════════════════════════
#  EXPORT
# ═══════════════════════════════════════════════════════════════
def _get_dec_data(slug):
    conn = get_db()
    dc = conn.execute("SELECT * FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc:
        conn.close(); return None
    is_author = session.get('user_id') == dc['user_id']
    alts = conn.execute("SELECT * FROM decision_alternatives WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
    crits = conn.execute("SELECT * FROM decision_criteria WHERE decision_id=? ORDER BY sort_order", (dc['id'],)).fetchall()
    rc = conn.execute("SELECT COUNT(*) as c FROM decision_responses WHERE decision_id=?", (dc['id'],)).fetchone()['c']
    ranking = []
    details = {}
    for a in alts:
        total = 0
        for c in crits:
            avg_row = conn.execute("SELECT AVG(score) as avg FROM decision_scores WHERE alternative_id=? AND criterion_id=?", (a['id'], c['id'])).fetchone()
            avg = round(avg_row['avg'], 2) if avg_row['avg'] else 0
            details[(a['id'], c['id'])] = avg
            total += avg
        avg_total = round(total / len(crits), 2) if crits else 0
        ranking.append({"id": a['id'], "name": a['name'], "score": avg_total})
    ranking.sort(key=lambda x: x['score'], reverse=True)
    respondents = []
    if is_author and dc['show_respondents'] and not dc['anonymous']:
        rrows = conn.execute("""SELECT u.first_name, u.last_name FROM decision_responses dr
            JOIN users u ON dr.user_id = u.id WHERE dr.decision_id=?""", (dc['id'],)).fetchall()
        respondents = [r['first_name'] + ' ' + r['last_name'] for r in rrows]
    conn.close()
    return {"title": dc['title'], "responses": rc, "ranking": ranking, "details": details,
            "alts": alts, "crits": crits, "respondents": respondents}


def _get_poll_data(slug):
    conn = get_db()
    p = conn.execute("SELECT * FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p:
        conn.close(); return None
    is_author = session.get('user_id') == p['user_id']
    opts = conn.execute("SELECT * FROM poll_options WHERE poll_id=? ORDER BY id", (p['id'],)).fetchall()
    vc = conn.execute("SELECT COUNT(DISTINCT COALESCE(user_id, fingerprint)) as c FROM poll_votes WHERE poll_id=?", (p['id'],)).fetchone()['c']
    results = []
    for o in opts:
        cnt = conn.execute("SELECT COUNT(*) as c FROM poll_votes WHERE option_id=?", (o['id'],)).fetchone()['c']
        voters = []
        if is_author and p['show_voters'] and not p['anonymous']:
            vrows = conn.execute("""SELECT u.first_name, u.last_name FROM poll_votes pv
                JOIN users u ON pv.user_id = u.id WHERE pv.option_id=?""", (o['id'],)).fetchall()
            voters = [r['first_name'] + ' ' + r['last_name'] for r in vrows]
        results.append({"text": o['text'], "votes": cnt, "voters": voters})
    conn.close()
    return {"title": p['title'], "totalVoters": vc, "results": results}


def _decision_winner_text(ranking):
    """Return 'Лучший вариант' string, handling ties."""
    if not ranking:
        return ""
    best_score = ranking[0]['score']
    winners = [r for r in ranking if abs(r['score'] - best_score) < 0.01]
    if len(winners) == 1:
        return f"Лучший вариант — {winners[0]['name']} ({best_score:.2f}/5)"
    names = ", ".join(w['name'] for w in winners)
    return f"Лучшие варианты (равный балл {best_score:.2f}/5) — {names}"


def _poll_winner_text(results):
    """Return 'Лидер голосования' string, handling ties."""
    if not results:
        return ""
    best_votes = max(r['votes'] for r in results)
    if best_votes == 0:
        return ""
    winners = [r for r in results if r['votes'] == best_votes]
    if len(winners) == 1:
        return f"Лидер голосования — {winners[0]['text']} ({best_votes} гол.)"
    names = ", ".join(w['text'] for w in winners)
    return f"Лидеры голосования (равный результат: {best_votes} гол.) — {names}"


def _build_review_context(d):
    """Build context string from decision data for AI prompts."""
    ranking = d['ranking']
    context = f"Вопрос: {d['title']}. Ответов: {d['responses']}. "
    context += "Рейтинг: " + ", ".join([f"{r['name']} ({r['score']:.2f}/5)" for r in ranking]) + ". "
    context += "Критерии: " + ", ".join([c['name'] for c in d['crits']]) + "."
    for r in ranking[:3]:
        parts = [f"{c['name']}: {d['details'].get((r['id'], c['id']), 0):.1f}" for c in d['crits']]
        context += f" {r['name']}: {', '.join(parts)}."
    return context


def _build_review_prompt(d, winner_line, is_tie, winners, best_score):
    """Build system prompt for AI review."""
    tie_instruction = ""
    if is_tie:
        tie_names = ", ".join(w['name'] for w in winners)
        tie_instruction = (
            f"ВНИМАНИЕ: варианты {tie_names} набрали одинаковый балл {best_score:.2f}/5! "
            "Обязательно отрази это, подробно сравни их между собой по каждому критерию "
            "и порекомендуй, какой из них всё же стоит выбрать и почему. "
        )
    return (
        "Ты эксперт-аналитик с глубокими знаниями в разных областях. "
        "Напиши развёрнутую, содержательную рецензию (8-12 предложений). "
        f"Начни с вывода: «{winner_line}». "
        + tie_instruction +
        "ВАЖНО — рецензия НЕ должна быть очевидной пересказкой баллов! Вместо этого:\n"
        "1) Объясни, ПОЧЕМУ именно этот вариант лидирует — раскрой причины, а не просто цифры.\n"
        "2) Приведи РЕАЛЬНЫЕ ФАКТЫ из своих знаний: конкретные цены, статистику, рейтинги, индексы, "
        "средние зарплаты, стоимость жизни, характеристики товаров — всё, что относится к теме вопроса. "
        "Например: «Стоимость жизни в Грузии по Numbeo — $600/мес., против $1200 в Чехии». "
        "Если тема касается товаров — реальные цены, ТТХ, отзывы.\n"
        "3) Укажи неочевидные подводные камни или преимущества, о которых пользователь мог не подумать.\n"
        "4) Дай чёткую, аргументированную итоговую рекомендацию.\n"
        "Отвечай на русском. Пиши живым языком, не канцеляритом."
    )


def _build_review_text(d):
    """Generate review text for export."""
    ranking = d['ranking']
    if not ranking or d['responses'] == 0:
        return ""
    winner = ranking[0]
    second = ranking[1] if len(ranking) > 1 else {"name": "—", "score": 0}
    gap = winner['score'] - second['score']
    best_score = winner['score']
    winners = [r for r in ranking if abs(r['score'] - best_score) < 0.01]
    is_tie = len(winners) > 1
    winner_line = _decision_winner_text(ranking)

    context = _build_review_context(d)
    sys_prompt = _build_review_prompt(d, winner_line, is_tie, winners, best_score)

    text, _, _ = _call_ai(sys_prompt, context, temperature=0.8, max_tokens=900)
    if text:
        return text

    fmt_data = {"winner": winner['name'], "score": winner['score'],
                "second": second['name'], "second_score": second['score'], "gap": gap}
    if is_tie:
        tpl = _REVIEW_TEMPLATES["tie"]
    elif winner['score'] < 2.5:
        tpl = _REVIEW_TEMPLATES["low"]
    elif gap < 0.5:
        tpl = _REVIEW_TEMPLATES["close"]
    else:
        tpl = _REVIEW_TEMPLATES["strong"]
    review = winner_line + "\n\n" + tpl.format(**fmt_data)
    strengths = [c['name'] for c in d['crits'] if d['details'].get((winner['id'], c['id']), 0) >= 4.0]
    weaknesses = [c['name'] for c in d['crits'] if d['details'].get((winner['id'], c['id']), 0) < 3.0]
    if strengths:
        review += f" Сильные стороны «{winner['name']}»: {', '.join(strengths)}."
    if weaknesses:
        review += f" Обратите внимание на: {', '.join(weaknesses)}."
    if is_tie:
        for w in winners[1:]:
            w_str = [c['name'] for c in d['crits'] if d['details'].get((w['id'], c['id']), 0) >= 4.0]
            w_weak = [c['name'] for c in d['crits'] if d['details'].get((w['id'], c['id']), 0) < 3.0]
            if w_str:
                review += f" Сильные стороны «{w['name']}»: {', '.join(w_str)}."
            if w_weak:
                review += f" Слабые стороны «{w['name']}»: {', '.join(w_weak)}."
    return review


@app.route('/api/decisions/<slug>/export/<fmt>')
def export_decision(slug, fmt):
    d = _get_dec_data(slug)
    if not d:
        return jsonify({"error": "Не найдено"}), 404
    detailed = request.args.get('mode') == 'detailed'
    review_text = _build_review_text(d) if detailed else ""

    winner_line = _decision_winner_text(d['ranking'])

    if fmt == 'txt':
        lines = ["Результаты: " + d['title'], "Ответов: " + str(d['responses']), ""]
        if winner_line:
            lines.append("★ " + winner_line)
            lines.append("")
        lines.append("Рейтинг:")
        for i, r in enumerate(d['ranking']):
            lines.append(f"  {i+1}. {r['name']} — {r['score']:.2f}/5")
        lines.append("")
        lines.append("Детализация по критериям:")
        header = [""] + [c['name'] for c in d['crits']]
        lines.append("\t".join(header))
        for r in d['ranking']:
            row = [r['name']]
            for c in d['crits']:
                row.append(f"{d['details'].get((r['id'], c['id']), 0):.1f}")
            lines.append("\t".join(row))
        if d.get('respondents'):
            lines += ["", "Участники: " + ", ".join(d['respondents'])]
        if review_text:
            lines += ["", "=" * 40, "РЕЦЕНЗИЯ:", "", review_text]
        buf = io.BytesIO("\n".join(lines).encode('utf-8'))
        from flask import send_file
        return send_file(buf, mimetype='text/plain', as_attachment=True, download_name=f'results-{slug}.txt')

    elif fmt == 'xlsx':
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        wb = Workbook()
        ws = wb.active
        ws.title = "Результаты"
        ws.append([d['title']])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(d['crits'])+2)
        ws['A1'].font = Font(bold=True, size=14)
        ws.append([f"Ответов: {d['responses']}"])
        if winner_line:
            ws.append([winner_line])
            ws[ws.max_row][0].font = Font(bold=True, size=12, color="2563EB")
        ws.append([])
        ws.append(["Место", "Вариант", "Общий балл"] + [c['name'] for c in d['crits']])
        hdr_fill = PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid")
        for cell in ws[4]:
            cell.font = Font(bold=True)
            cell.fill = hdr_fill
        for i, r in enumerate(d['ranking']):
            row = [i+1, r['name'], r['score']]
            for c in d['crits']:
                row.append(d['details'].get((r['id'], c['id']), 0))
            ws.append(row)
        for col in ws.columns:
            try:
                letter = col[0].column_letter
            except AttributeError:
                continue
            max_len = max(len(str(cell.value or '')) for cell in col if hasattr(cell, 'column_letter'))
            if max_len:
                ws.column_dimensions[letter].width = min(max_len + 4, 30)
        if d.get('respondents'):
            ws.append([]); ws.append(["Участники:"])
            ws[ws.max_row][0].font = Font(bold=True)
            ws.append([", ".join(d['respondents'])])
        if review_text:
            ws.append([]); ws.append(["Рецензия:"])
            ws.append([review_text])
            ws[ws.max_row][0].font = Font(italic=True)
        buf = io.BytesIO()
        wb.save(buf); buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True, download_name=f'results-{slug}.xlsx')

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        doc.add_heading(d['title'], level=1)
        doc.add_paragraph(f"Ответов: {d['responses']}")
        if winner_line:
            wp = doc.add_paragraph()
            wr = wp.add_run(winner_line)
            wr.bold = True
            wr.font.size = Pt(13)
            wr.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        doc.add_heading("Рейтинг", level=2)
        for i, r in enumerate(d['ranking']):
            p = doc.add_paragraph()
            run = p.add_run(f"{i+1}. {r['name']}")
            run.bold = True
            p.add_run(f" — {r['score']:.2f} из 5")
        doc.add_heading("Детализация", level=2)
        cols = len(d['crits']) + 1
        table = doc.add_table(rows=1, cols=cols, style='Light Grid Accent 1')
        hdr = table.rows[0].cells
        hdr[0].text = "Вариант"
        for j, c in enumerate(d['crits']):
            hdr[j+1].text = c['name']
        for r in d['ranking']:
            row = table.add_row().cells
            row[0].text = r['name']
            for j, c in enumerate(d['crits']):
                row[j+1].text = f"{d['details'].get((r['id'], c['id']), 0):.1f}"
        if d.get('respondents'):
            doc.add_heading("Участники", level=2)
            doc.add_paragraph(", ".join(d['respondents']))
        if review_text:
            doc.add_heading("Рецензия", level=2)
            rp = doc.add_paragraph(review_text)
            rp.style.font.size = Pt(11)
        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=f'results-{slug}.docx')

    elif fmt == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        _FONT = 'SegoeUI'
        if _FONT not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(_FONT, r'C:\Windows\Fonts\segoeui.ttf'))
            pdfmetrics.registerFont(TTFont(_FONT + '-Bold', r'C:\Windows\Fonts\segoeuib.ttf'))
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            registerFontFamily(_FONT, normal=_FONT, bold=_FONT + '-Bold')

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        for s in styles.byName.values():
            s.fontName = _FONT
        elements = []

        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontName=_FONT, fontSize=16, spaceAfter=6)
        elements.append(Paragraph(d['title'], title_style))
        elements.append(Paragraph(f"Ответов: {d['responses']}", styles['Normal']))
        if winner_line:
            winner_style = ParagraphStyle('WinnerLine', fontName=_FONT + '-Bold', fontSize=13, textColor=colors.HexColor('#2563EB'), spaceBefore=8, spaceAfter=4)
            elements.append(Paragraph(winner_line, winner_style))
        elements.append(Spacer(1, 12))

        elements.append(Paragraph("Рейтинг", styles['Heading2']))
        for i, r in enumerate(d['ranking']):
            elements.append(Paragraph(f"{i+1}. <b>{r['name']}</b> — {r['score']:.2f} из 5", styles['Normal']))
        elements.append(Spacer(1, 12))

        elements.append(Paragraph("Детализация по критериям", styles['Heading2']))
        header = ["Вариант"] + [c['name'] for c in d['crits']]
        table_data = [header]
        for r in d['ranking']:
            row = [r['name']]
            for c in d['crits']:
                row.append(f"{d['details'].get((r['id'], c['id']), 0):.1f}")
            table_data.append(row)
        t = Table(table_data, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#DBEAFE')),
            ('FONTNAME', (0,0), (-1,0), _FONT + '-Bold'),
            ('FONTNAME', (0,1), (-1,-1), _FONT),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
            ('ALIGN', (1,1), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
        ]))
        elements.append(t)
        if d.get('respondents'):
            elements.append(Spacer(1, 12))
            elements.append(Paragraph("Участники", styles['Heading2']))
            elements.append(Paragraph(", ".join(d['respondents']), styles['Normal']))
        if review_text:
            elements.append(Spacer(1, 18))
            elements.append(Paragraph("Рецензия", styles['Heading2']))
            elements.append(Paragraph(review_text, styles['Normal']))
        doc.build(elements)
        buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name=f'results-{slug}.pdf')

    return jsonify({"error": "Неизвестный формат"}), 400


@app.route('/api/polls/<slug>/export/<fmt>')
def export_poll(slug, fmt):
    d = _get_poll_data(slug)
    if not d:
        return jsonify({"error": "Не найдено"}), 404
    total = sum(r['votes'] for r in d['results']) or 1

    poll_winner = _poll_winner_text(d['results'])

    if fmt == 'txt':
        lines = ["Результаты голосования: " + d['title'], f"Проголосовало: {d['totalVoters']}", ""]
        if poll_winner:
            lines.append("★ " + poll_winner)
            lines.append("")
        for r in d['results']:
            pct = round(r['votes'] / total * 100)
            line = f"  {r['text']} — {r['votes']} голосов ({pct}%)"
            if r.get('voters'):
                line += f"  [Голосовали: {', '.join(r['voters'])}]"
            lines.append(line)
        buf = io.BytesIO("\n".join(lines).encode('utf-8'))
        from flask import send_file
        return send_file(buf, mimetype='text/plain', as_attachment=True, download_name=f'poll-{slug}.txt')

    elif fmt == 'xlsx':
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        wb = Workbook(); ws = wb.active; ws.title = "Результаты"
        ws.append([d['title']]); ws['A1'].font = Font(bold=True, size=14)
        ws.append([f"Проголосовало: {d['totalVoters']}"])
        if poll_winner:
            ws.append([poll_winner])
            ws[ws.max_row][0].font = Font(bold=True, size=12, color="2563EB")
        ws.append([])
        has_voters = any(r.get('voters') for r in d['results'])
        header = ["Вариант", "Голосов", "%"]
        if has_voters:
            header.append("Голосовавшие")
        ws.append(header)
        hdr_row = ws.max_row
        for cell in ws[hdr_row]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid")
        for r in d['results']:
            row = [r['text'], r['votes'], round(r['votes']/total*100)]
            if has_voters:
                row.append(", ".join(r.get('voters', [])))
            ws.append(row)
        for col in ws.columns:
            try:
                letter = col[0].column_letter
            except AttributeError:
                continue
            max_len = max(len(str(cell.value or '')) for cell in col if hasattr(cell, 'column_letter'))
            if max_len:
                ws.column_dimensions[letter].width = min(max_len + 4, 30)
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True, download_name=f'poll-{slug}.xlsx')

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        doc.add_heading(d['title'], level=1)
        doc.add_paragraph(f"Проголосовало: {d['totalVoters']}")
        if poll_winner:
            wp = doc.add_paragraph()
            wr = wp.add_run(poll_winner)
            wr.bold = True
            wr.font.size = Pt(13)
            wr.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        has_voters = any(r.get('voters') for r in d['results'])
        ncols = 4 if has_voters else 3
        table = doc.add_table(rows=1, cols=ncols, style='Light Grid Accent 1')
        hdr = table.rows[0].cells
        hdr[0].text = "Вариант"; hdr[1].text = "Голосов"; hdr[2].text = "%"
        if has_voters:
            hdr[3].text = "Голосовавшие"
        for r in d['results']:
            row = table.add_row().cells
            row[0].text = r['text']; row[1].text = str(r['votes']); row[2].text = f"{round(r['votes']/total*100)}%"
            if has_voters:
                row[3].text = ", ".join(r.get('voters', []))
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=f'poll-{slug}.docx')

    elif fmt == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        _FONT = 'SegoeUI'
        if _FONT not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(_FONT, r'C:\Windows\Fonts\segoeui.ttf'))
            pdfmetrics.registerFont(TTFont(_FONT + '-Bold', r'C:\Windows\Fonts\segoeuib.ttf'))
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            registerFontFamily(_FONT, normal=_FONT, bold=_FONT + '-Bold')

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        for s in styles.byName.values():
            s.fontName = _FONT
        elements = []
        elements.append(Paragraph(d['title'], ParagraphStyle('T', parent=styles['Title'], fontName=_FONT, fontSize=16, spaceAfter=6)))
        elements.append(Paragraph(f"Проголосовало: {d['totalVoters']}", styles['Normal']))
        if poll_winner:
            pw_style = ParagraphStyle('PollWinner', fontName=_FONT + '-Bold', fontSize=13, textColor=colors.HexColor('#2563EB'), spaceBefore=8, spaceAfter=4)
            elements.append(Paragraph(poll_winner, pw_style))
        elements.append(Spacer(1, 12))
        has_voters = any(r.get('voters') for r in d['results'])
        header = ["Вариант", "Голосов", "%"]
        if has_voters:
            header.append("Голосовавшие")
        table_data = [header]
        for r in d['results']:
            row = [r['text'], str(r['votes']), f"{round(r['votes']/total*100)}%"]
            if has_voters:
                row.append(", ".join(r.get('voters', [])))
            table_data.append(row)
        t = Table(table_data, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#DBEAFE')),
            ('FONTNAME', (0,0), (-1,0), _FONT + '-Bold'),
            ('FONTNAME', (0,1), (-1,-1), _FONT),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
            ('ALIGN', (1,0), (-1,-1), 'CENTER'),
        ]))
        elements.append(t); doc.build(elements); buf.seek(0)
        from flask import send_file
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name=f'poll-{slug}.pdf')

    return jsonify({"error": "Неизвестный формат"}), 400


# ═══════════════════════════════════════════════════════════════
#  AI REVIEW
# ═══════════════════════════════════════════════════════════════
_REVIEW_TEMPLATES = {
    "strong": "Лучший вариант — «{winner}» с баллом {score:.2f}/5. Результаты достаточно однозначны: лидер уверенно опережает ближайшего конкурента «{second}» на {gap:.2f} балла. Это явное предпочтение респондентов. Рекомендую остановиться на варианте «{winner}», так как он показал превосходство по большинству критериев.",
    "close": "Лучший вариант — «{winner}» ({score:.2f}/5), однако результаты плотные: «{second}» отстаёт всего на {gap:.2f} балла. При таком небольшом разрыве стоит обратить внимание на отдельные критерии — возможно, по ключевому для вас параметру второй вариант даже выигрывает. Тем не менее, по совокупности оценок рекомендую «{winner}».",
    "tie": "Результаты практически равные — «{winner}» и «{second}» набрали почти одинаковый балл ({score:.2f} vs {second_score:.2f}). В такой ситуации решение не может быть основано только на цифрах. Рекомендую внимательно сравнить их по отдельным критериям: тот вариант, который выигрывает по наиболее важным для вас параметрам — и есть лучший выбор.",
    "low": "Лучший вариант — «{winner}», но его балл всего {score:.2f}/5, что довольно невысоко. Это может означать, что ни один из предложенных вариантов не является идеальным. Тем не менее, «{winner}» — лучший из имеющихся. Рекомендую рассмотреть дополнительные альтернативы или скорректировать критерии.",
}

@app.route('/api/decisions/<slug>/ai-review')
def ai_review(slug):
    d = _get_dec_data(slug)
    if not d or d['responses'] == 0:
        return jsonify({"error": "Нет данных"}), 400

    ranking = d['ranking']
    winner = ranking[0]
    second = ranking[1] if len(ranking) > 1 else {"name": "—", "score": 0}
    gap = winner['score'] - second['score']
    best_score = winner['score']
    winners = [r for r in ranking if abs(r['score'] - best_score) < 0.01]
    is_tie = len(winners) > 1
    winner_line = _decision_winner_text(ranking)

    context = _build_review_context(d)
    sys_prompt = _build_review_prompt(d, winner_line, is_tie, winners, best_score)

    text, source, ai_error = _call_ai(sys_prompt, context, temperature=0.8, max_tokens=900)
    if text:
        return jsonify({"review": text, "source": source})

    fmt_data = {"winner": winner['name'], "score": winner['score'],
                "second": second['name'], "second_score": second['score'], "gap": gap}
    if is_tie:
        tpl = _REVIEW_TEMPLATES["tie"]
    elif winner['score'] < 2.5:
        tpl = _REVIEW_TEMPLATES["low"]
    elif gap < 0.5:
        tpl = _REVIEW_TEMPLATES["close"]
    else:
        tpl = _REVIEW_TEMPLATES["strong"]

    review = winner_line + "\n\n" + tpl.format(**fmt_data)

    strengths = []
    weaknesses = []
    for c in d['crits']:
        val = d['details'].get((winner['id'], c['id']), 0)
        if val >= 4.0:
            strengths.append(c['name'])
        elif val < 3.0:
            weaknesses.append(c['name'])
    if strengths:
        review += f" Сильные стороны «{winner['name']}»: {', '.join(strengths)}."
    if weaknesses:
        review += f" Обратите внимание на: {', '.join(weaknesses)} — здесь оценки ниже среднего."
    if is_tie:
        for w in winners[1:]:
            w_str = [c['name'] for c in d['crits'] if d['details'].get((w['id'], c['id']), 0) >= 4.0]
            w_weak = [c['name'] for c in d['crits'] if d['details'].get((w['id'], c['id']), 0) < 3.0]
            if w_str:
                review += f" Сильные стороны «{w['name']}»: {', '.join(w_str)}."
            if w_weak:
                review += f" Слабые стороны «{w['name']}»: {', '.join(w_weak)}."
        review += " Рекомендую сравнить эти варианты по наиболее важному для вас критерию и выбрать тот, который выигрывает именно по нему."

    return jsonify({"review": review, "source": "local", "ai_error": ai_error})


# ═══════════════════════════════════════════════════════════════
#  PUBLIC FEED, REACTIONS, COMMENTS
# ═══════════════════════════════════════════════════════════════
@app.route('/api/feed')
def public_feed():
    @after_this_request
    def _no_cache(response):
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
        return response
    try:
        page = max(int(request.args.get('page', 1)), 1)
        per_page = 10
        offset = (page - 1) * per_page
        uid = session.get('user_id')
        conn = get_db()
        combined = []
        try:
            conn.execute("SELECT is_public FROM polls LIMIT 1")
        except sqlite3.OperationalError:
            try:
                conn.execute("ALTER TABLE polls ADD COLUMN is_public INTEGER DEFAULT 0")
                conn.execute("ALTER TABLE decisions ADD COLUMN is_public INTEGER DEFAULT 0")
                conn.commit()
            except sqlite3.OperationalError:
                pass
        polls = conn.execute(
            "SELECT p.*, u.first_name, u.last_name FROM polls p JOIN users u ON p.user_id=u.id "
            "WHERE COALESCE(p.is_public, 0) = 1 ORDER BY p.created_at DESC", ()).fetchall()
        print('[Feed] raw polls=%d' % len(polls))
        for p in polls:
            pid = p['id']
            tv = conn.execute("SELECT COUNT(DISTINCT COALESCE(user_id,fingerprint)) as c FROM poll_votes WHERE poll_id=?", (pid,)).fetchone()['c']
            has_voted = False
            if uid:
                has_voted = conn.execute("SELECT 1 FROM poll_votes WHERE poll_id=? AND user_id=?", (pid, uid)).fetchone() is not None
            likes = conn.execute("SELECT COUNT(*) as c FROM poll_reactions WHERE poll_id=? AND type='like'", (pid,)).fetchone()['c']
            dislikes = conn.execute("SELECT COUNT(*) as c FROM poll_reactions WHERE poll_id=? AND type='dislike'", (pid,)).fetchone()['c']
            comment_count = conn.execute("SELECT COUNT(*) as c FROM poll_comments WHERE poll_id=?", (pid,)).fetchone()['c']
            my_reaction = None
            if uid:
                r = conn.execute("SELECT type FROM poll_reactions WHERE poll_id=? AND user_id=?", (pid, uid)).fetchone()
                if r: my_reaction = r['type']
            last_comments = conn.execute(
                "SELECT c.text, c.created_at, u.first_name, u.last_name FROM poll_comments c "
                "JOIN users u ON c.user_id=u.id WHERE c.poll_id=? ORDER BY c.created_at DESC LIMIT 2", (pid,)).fetchall()
            active = bool(p['is_active']) and check_deadline(p['deadline'])
            combined.append({
                "type": "poll", "slug": p['slug'], "title": p['title'], "description": p['description'],
                "authorName": ((p['first_name'] or '') + ' ' + (p['last_name'] or '')).strip(),
                "authorInitials": ((p['first_name'] or ' ')[0] + (p['last_name'] or ' ')[0]).upper(),
                "createdAt": p['created_at'], "isActive": active,
                "totalVotes": tv, "likes": likes, "dislikes": dislikes,
                "commentCount": comment_count, "myReaction": my_reaction,
                "hasVoted": has_voted, "showResults": (p['show_results'] if p['show_results'] else 'always'),
                "lastComments": [{"text": c['text'], "author": (c['first_name'] or '') + ' ' + (c['last_name'] or ''),
                                  "createdAt": c['created_at']} for c in reversed(list(last_comments))]
            })

        decs = conn.execute(
            "SELECT d.*, u.first_name, u.last_name FROM decisions d JOIN users u ON d.user_id=u.id "
            "WHERE COALESCE(d.is_public, 0) = 1 ORDER BY d.created_at DESC", ()).fetchall()
        print('[Feed] raw decs=%d' % len(decs))
        for dc in decs:
            did = dc['id']
            has_responded = False
            if uid:
                has_responded = conn.execute("SELECT 1 FROM decision_responses WHERE decision_id=? AND user_id=?", (did, uid)).fetchone() is not None
            rc = conn.execute("SELECT COUNT(*) as c FROM decision_responses WHERE decision_id=?", (did,)).fetchone()['c']
            ac = conn.execute("SELECT COUNT(*) as c FROM decision_alternatives WHERE decision_id=?", (did,)).fetchone()['c']
            cc = conn.execute("SELECT COUNT(*) as c FROM decision_criteria WHERE decision_id=?", (did,)).fetchone()['c']
            comment_count = conn.execute("SELECT COUNT(*) as c FROM decision_comments WHERE decision_id=?", (did,)).fetchone()['c']
            likes = conn.execute("SELECT COUNT(*) as c FROM decision_reactions WHERE decision_id=? AND type='like'", (did,)).fetchone()['c']
            dislikes = conn.execute("SELECT COUNT(*) as c FROM decision_reactions WHERE decision_id=? AND type='dislike'", (did,)).fetchone()['c']
            my_reaction = None
            if uid:
                r = conn.execute("SELECT type FROM decision_reactions WHERE decision_id=? AND user_id=?", (did, uid)).fetchone()
                if r: my_reaction = r['type']
            last_comments = conn.execute(
                "SELECT c.text, c.created_at, u.first_name, u.last_name FROM decision_comments c "
                "JOIN users u ON c.user_id=u.id WHERE c.decision_id=? ORDER BY c.created_at DESC LIMIT 2", (did,)).fetchall()
            active = bool(dc['is_active']) and check_deadline(dc['deadline'])
            sr = (dc['show_results'] if dc['show_results'] else 'always')
            combined.append({
                "type": "decision", "slug": dc['slug'], "title": dc['title'], "description": dc['description'],
                "authorName": ((dc['first_name'] or '') + ' ' + (dc['last_name'] or '')).strip(),
                "authorInitials": ((dc['first_name'] or ' ')[0] + (dc['last_name'] or ' ')[0]).upper(),
                "createdAt": dc['created_at'], "isActive": active,
                "responsesCount": rc, "alternativesCount": ac, "criteriaCount": cc,
                "likes": likes, "dislikes": dislikes, "commentCount": comment_count, "myReaction": my_reaction,
                "hasResponded": has_responded, "showResults": sr,
                "lastComments": [{"text": c['text'], "author": (c['first_name'] or '') + ' ' + (c['last_name'] or ''),
                                 "createdAt": c['created_at']} for c in reversed(list(last_comments))]
            })

        conn.close()
        combined.sort(key=lambda x: x.get('createdAt') or '', reverse=True)
        total = len(combined)
        items = combined[offset:offset + per_page]
        print('[Feed] polls=%d decs=%d total=%d items=%d' % (len(polls), len(decs), total, len(items)))
        return jsonify({"items": items, "total": total, "page": page,
                        "pages": max(1, (total + per_page - 1) // per_page)})
    except Exception as e:
        import traceback
        print('[Feed error]', e)
        traceback.print_exc()
        try:
            conn.close()
        except Exception:
            pass
        return jsonify({"items": [], "total": 0, "page": 1, "pages": 1, "error": str(e)})

@app.route('/api/polls/<slug>/react', methods=['POST'])
@block_if_banned
def react_poll(slug):
    d = request.get_json() or {}
    rtype = d.get('type', '')
    if rtype not in ('like', 'dislike'):
        return jsonify({"error": "Invalid reaction type"}), 400
    conn = get_db()
    p = conn.execute("SELECT id FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    uid = session['user_id']
    existing = conn.execute("SELECT id, type FROM poll_reactions WHERE poll_id=? AND user_id=?", (p['id'], uid)).fetchone()
    if existing:
        if existing['type'] == rtype:
            conn.execute("DELETE FROM poll_reactions WHERE id=?", (existing['id'],))
        else:
            conn.execute("UPDATE poll_reactions SET type=? WHERE id=?", (rtype, existing['id']))
    else:
        conn.execute("INSERT INTO poll_reactions (poll_id,user_id,type) VALUES (?,?,?)", (p['id'], uid, rtype))
    conn.commit()
    likes = conn.execute("SELECT COUNT(*) as c FROM poll_reactions WHERE poll_id=? AND type='like'", (p['id'],)).fetchone()['c']
    dislikes = conn.execute("SELECT COUNT(*) as c FROM poll_reactions WHERE poll_id=? AND type='dislike'", (p['id'],)).fetchone()['c']
    my = conn.execute("SELECT type FROM poll_reactions WHERE poll_id=? AND user_id=?", (p['id'], uid)).fetchone()
    conn.close()
    return jsonify({"ok": True, "likes": likes, "dislikes": dislikes, "myReaction": my['type'] if my else None})


@app.route('/api/decisions/<slug>/react', methods=['POST'])
@block_if_banned
def react_decision(slug):
    d = request.get_json() or {}
    rtype = d.get('type', '')
    if rtype not in ('like', 'dislike'):
        return jsonify({"error": "Invalid reaction type"}), 400
    conn = get_db()
    dc = conn.execute("SELECT id FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    uid = session['user_id']
    existing = conn.execute("SELECT id, type FROM decision_reactions WHERE decision_id=? AND user_id=?", (dc['id'], uid)).fetchone()
    if existing:
        if existing['type'] == rtype:
            conn.execute("DELETE FROM decision_reactions WHERE id=?", (existing['id'],))
        else:
            conn.execute("UPDATE decision_reactions SET type=? WHERE id=?", (rtype, existing['id']))
    else:
        conn.execute("INSERT INTO decision_reactions (decision_id,user_id,type) VALUES (?,?,?)", (dc['id'], uid, rtype))
    conn.commit()
    likes = conn.execute("SELECT COUNT(*) as c FROM decision_reactions WHERE decision_id=? AND type='like'", (dc['id'],)).fetchone()['c']
    dislikes = conn.execute("SELECT COUNT(*) as c FROM decision_reactions WHERE decision_id=? AND type='dislike'", (dc['id'],)).fetchone()['c']
    my = conn.execute("SELECT type FROM decision_reactions WHERE decision_id=? AND user_id=?", (dc['id'], uid)).fetchone()
    conn.close()
    return jsonify({"ok": True, "likes": likes, "dislikes": dislikes, "myReaction": my['type'] if my else None})


@app.route('/api/polls/<slug>/comments')
def poll_comments_list(slug):
    conn = get_db()
    p = conn.execute("SELECT id FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    rows = conn.execute(
        "SELECT c.id, c.text, c.created_at, u.first_name, u.last_name FROM poll_comments c "
        "JOIN users u ON c.user_id=u.id WHERE c.poll_id=? ORDER BY c.created_at ASC", (p['id'],)).fetchall()
    conn.close()
    return jsonify({"comments": [{"id": r['id'], "text": r['text'], "author": r['first_name'] + ' ' + r['last_name'],
                                   "createdAt": r['created_at']} for r in rows]})

@app.route('/api/polls/<slug>/comments', methods=['POST'])
@block_if_banned
def add_poll_comment(slug):
    d = request.get_json() or {}
    text = (d.get('text') or '').strip()
    if not text: return jsonify({"error": "Введите комментарий"}), 400
    if len(text) > 1000: return jsonify({"error": "Максимум 1000 символов"}), 400
    ok, reason = _moderate_content(text)
    if not ok:
        w, blocked = _apply_warning(session['user_id'], 'comment', text, '', reason)
        resp = {"error": "Комментарий отклонён модерацией", "warning": True,
                "reason": reason, "warningCount": w}
        if blocked:
            resp["blocked"] = True
            resp["error"] = "Аккаунт заблокирован за повторные нарушения"
        return jsonify(resp), 403
    conn = get_db()
    p = conn.execute("SELECT id FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("INSERT INTO poll_comments (poll_id,user_id,text) VALUES (?,?,?)", (p['id'], session['user_id'], text))
    conn.commit()
    user = conn.execute("SELECT first_name, last_name FROM users WHERE id=?", (session['user_id'],)).fetchone()
    conn.close()
    return jsonify({"ok": True, "author": user['first_name'] + ' ' + user['last_name']}), 201


@app.route('/api/decisions/<slug>/comments')
def decision_comments_list(slug):
    conn = get_db()
    dc = conn.execute("SELECT id FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    rows = conn.execute(
        "SELECT c.id, c.text, c.created_at, u.first_name, u.last_name FROM decision_comments c "
        "JOIN users u ON c.user_id=u.id WHERE c.decision_id=? ORDER BY c.created_at ASC", (dc['id'],)).fetchall()
    conn.close()
    return jsonify({"comments": [{"id": r['id'], "text": r['text'], "author": r['first_name'] + ' ' + r['last_name'],
                                 "createdAt": r['created_at']} for r in rows]})


@app.route('/api/decisions/<slug>/comments', methods=['POST'])
@block_if_banned
def add_decision_comment(slug):
    d = request.get_json() or {}
    text = (d.get('text') or '').strip()
    if not text: return jsonify({"error": "Введите комментарий"}), 400
    if len(text) > 1000: return jsonify({"error": "Максимум 1000 символов"}), 400
    ok, reason = _moderate_content(text)
    if not ok:
        w, blocked = _apply_warning(session['user_id'], 'comment', text, '', reason)
        resp = {"error": "Комментарий отклонён модерацией", "warning": True,
                "reason": reason, "warningCount": w}
        if blocked:
            resp["blocked"] = True
            resp["error"] = "Аккаунт заблокирован за повторные нарушения"
        return jsonify(resp), 403
    conn = get_db()
    dc = conn.execute("SELECT id FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc: conn.close(); return jsonify({"error": "Не найдено"}), 404
    conn.execute("INSERT INTO decision_comments (decision_id,user_id,text) VALUES (?,?,?)", (dc['id'], session['user_id'], text))
    conn.commit()
    user = conn.execute("SELECT first_name, last_name FROM users WHERE id=?", (session['user_id'],)).fetchone()
    conn.close()
    return jsonify({"ok": True, "author": user['first_name'] + ' ' + user['last_name']}), 201


# ═══════════════════════════════════════════════════════════════
#  ADMIN
# ═══════════════════════════════════════════════════════════════
@app.route('/api/admin/users')
@admin_required
def admin_list_users():
    conn = get_db()
    rows = conn.execute(
        "SELECT id, email, first_name, last_name, created_at, warnings, is_blocked FROM users ORDER BY id"
    ).fetchall()
    conn.close()
    return jsonify({"users": [{"id": r["id"], "email": r["email"], "firstName": r["first_name"],
                               "lastName": r["last_name"], "createdAt": r["created_at"],
                               "warnings": r["warnings"], "isBlocked": bool(r["is_blocked"])} for r in rows]})


@app.route('/api/admin/users/<int:user_id>/unban', methods=['POST'])
@admin_required
def admin_unban_user(user_id):
    conn = get_db()
    u = conn.execute("SELECT id FROM users WHERE id=?", (user_id,)).fetchone()
    if not u:
        conn.close()
        return jsonify({"error": "Пользователь не найден"}), 404
    conn.execute("UPDATE users SET is_blocked=0, warnings=0 WHERE id=?", (user_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route('/api/admin/moderation-log')
@admin_required
def admin_moderation_log():
    conn = get_db()
    rows = conn.execute(
        "SELECT m.id, m.user_id, m.content_type, m.poll_title, m.poll_description, m.reason, m.created_at, "
        "u.email, u.first_name, u.last_name FROM moderation_log m "
        "JOIN users u ON m.user_id=u.id ORDER BY m.id DESC LIMIT 200"
    ).fetchall()
    conn.close()
    return jsonify({"log": [{"id": r["id"], "userId": r["user_id"], "userEmail": r["email"],
                             "userName": (r["first_name"] or "") + " " + (r["last_name"] or ""),
                             "contentType": r["content_type"], "pollTitle": r["poll_title"],
                             "pollDescription": (r["poll_description"] or "")[:200],
                             "reason": r["reason"], "createdAt": r["created_at"]} for r in rows]})


@app.route('/api/admin/check')
@admin_required
def admin_check():
    return jsonify({"ok": True, "admin": True})


@app.route('/api/admin/polls')
@admin_required
def admin_polls():
    """Список всех опросов для модерации."""
    conn = get_db()
    rows = conn.execute(
        "SELECT p.id, p.slug, p.title, p.created_at, p.is_public, u.email, u.first_name, u.last_name "
        "FROM polls p JOIN users u ON p.user_id=u.id ORDER BY p.created_at DESC"
    ).fetchall()
    conn.close()
    return jsonify({"polls": [{"id": r["id"], "slug": r["slug"], "title": r["title"],
        "createdAt": r["created_at"], "isPublic": bool(r["is_public"]),
        "authorEmail": r["email"], "authorName": (r["first_name"] or "") + " " + (r["last_name"] or "")} for r in rows]})


@app.route('/api/admin/decisions')
@admin_required
def admin_decisions():
    """Список всех решений для модерации."""
    conn = get_db()
    rows = conn.execute(
        "SELECT d.id, d.slug, d.title, d.created_at, d.is_public, u.email, u.first_name, u.last_name "
        "FROM decisions d JOIN users u ON d.user_id=u.id ORDER BY d.created_at DESC"
    ).fetchall()
    conn.close()
    return jsonify({"decisions": [{"id": r["id"], "slug": r["slug"], "title": r["title"],
        "createdAt": r["created_at"], "isPublic": bool(r["is_public"]),
        "authorEmail": r["email"], "authorName": (r["first_name"] or "") + " " + (r["last_name"] or "")} for r in rows]})


@app.route('/api/admin/polls/<slug>', methods=['DELETE'])
@admin_required
def admin_delete_poll(slug):
    """Удалить опрос (модерация)."""
    conn = get_db()
    p = conn.execute("SELECT id FROM polls WHERE slug=?", (slug,)).fetchone()
    if not p:
        conn.close()
        return jsonify({"error": "Не найдено"}), 404
    conn.execute("DELETE FROM polls WHERE id=?", (p["id"],))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route('/api/admin/decisions/<slug>', methods=['DELETE'])
@admin_required
def admin_delete_decision(slug):
    """Удалить решение (модерация)."""
    conn = get_db()
    dc = conn.execute("SELECT id FROM decisions WHERE slug=?", (slug,)).fetchone()
    if not dc:
        conn.close()
        return jsonify({"error": "Не найдено"}), 404
    conn.execute("DELETE FROM decisions WHERE id=?", (dc["id"],))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route('/api/admin/publish-all-to-feed', methods=['POST'])
@admin_required
def admin_publish_all_to_feed():
    """Выставить is_public=1 у всех опросов и решений — чтобы они снова отображались в ленте."""
    conn = get_db()
    try:
        cols_p = [r['name'] for r in conn.execute("PRAGMA table_info(polls)").fetchall()]
        cols_d = [r['name'] for r in conn.execute("PRAGMA table_info(decisions)").fetchall()]
        if 'is_public' not in cols_p:
            conn.execute("ALTER TABLE polls ADD COLUMN is_public INTEGER DEFAULT 0")
        if 'is_public' not in cols_d:
            conn.execute("ALTER TABLE decisions ADD COLUMN is_public INTEGER DEFAULT 0")
        conn.execute("UPDATE polls SET is_public = 1")
        conn.execute("UPDATE decisions SET is_public = 1")
        conn.commit()
        np = conn.execute("SELECT COUNT(*) as c FROM polls WHERE is_public=1").fetchone()['c']
        nd = conn.execute("SELECT COUNT(*) as c FROM decisions WHERE is_public=1").fetchone()['c']
    except Exception as e:
        conn.close()
        print('[publish-all-to-feed]', e)
        return jsonify({"ok": False, "error": str(e)}), 500
    conn.close()
    return jsonify({"ok": True, "polls": np, "decisions": nd})


# ═══════════════════════════════════════════════════════════════
#  FEEDBACK (banned users)
# ═══════════════════════════════════════════════════════════════
@app.route('/api/feedback', methods=['POST'])
def send_feedback():
    d = request.get_json() or {}
    text = (d.get('text') or '').strip()
    email = (d.get('email') or '').strip()
    if not text:
        return jsonify({"error": "Введите сообщение"}), 400
    uid = session.get('user_id')
    if uid and not email:
        conn = get_db()
        user = conn.execute("SELECT email FROM users WHERE id=?", (uid,)).fetchone()
        conn.close()
        if user:
            email = user['email']
    body = f"""
    <h3>Обратная связь / Запрос разблокировки</h3>
    <p><b>Email пользователя:</b> {email or 'не указан'}</p>
    <p><b>Сообщение:</b></p>
    <p>{text}</p>
    """
    sent, err_msg = _send_email(SMTP_EMAIL, 'Обратная связь / Запрос разблокировки', body)
    if sent:
        return jsonify({"ok": True})
    return jsonify({"error": err_msg or "Не удалось отправить письмо. Попробуйте позже."}), 503


# ─── Статика (последним, чтобы не перехватывать /api/*) ───────
@app.route('/<path:path>', methods=['GET', 'HEAD', 'OPTIONS', 'POST'])
def serve_static(path):
    if path.startswith('api/'):
        return jsonify({"error": "Not found"}), 404
    if os.path.isfile(os.path.join('.', path)):
        return send_from_directory('.', path)
    return send_from_directory('.', 'index.html')


# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    import ssl
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(('8.8.8.8', 80))
        _lan_ip = s.getsockname()[0]
        s.close()
    except Exception:
        _lan_ip = None
    use_ssl = os.path.isfile('cert.pem') and os.path.isfile('key.pem')
    if use_ssl:
        ctx = ssl.SSLContext(ssl.PROTOCOL_TLS_SERVER)
        ctx.load_cert_chain('cert.pem', 'key.pem')
        print("Сервер: https://localhost:5000")
        if _lan_ip:
            print("  Для доступа по сети: https://%s:5000" % _lan_ip)
        app.run(host='0.0.0.0', debug=True, port=5000, ssl_context=ctx)
    else:
        print("Сервер: http://localhost:5000")
        if _lan_ip:
            print("  Для друга в той же сети: http://%s:5000" % _lan_ip)
        print("  (для HTTPS: openssl req -x509 -newkey rsa:2048 -keyout key.pem -out cert.pem -days 365 -nodes)")
        app.run(host='0.0.0.0', debug=True, port=5000)
